"""首板布局专用的历史模型知识库。

这个模块不导入现有 ``review_app.knowledge.KnowledgeStore``，也不读取
``review_knowledge.db``。首板布局的文件、分片、FTS 和向量索引全部保存在
``dragon_knowledge.db`` 中；FastEmbed 模型缓存可以复用本机已有的只读文件。
"""

from __future__ import annotations

from array import array
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import datetime
import csv
import io
import json
import math
import os
from pathlib import Path
import re
import sqlite3
import threading
from typing import Any, Iterable, Mapping, Sequence
from uuid import uuid4

try:  # 复用现有项目的中文清洗与分片逻辑；独立运行时保留轻量回退。
    from ..cleaning import content_hash as _project_content_hash
    from ..cleaning import normalize_text as _project_normalize_text
    from ..cleaning import split_chunks as _project_split_chunks
except ImportError:  # pragma: no cover - 仅用于单文件调试
    _project_content_hash = None
    _project_normalize_text = None
    _project_split_chunks = None

try:
    from ..config import DATA_DIR
except ImportError:  # pragma: no cover - 仅用于单文件调试
    DATA_DIR = Path(__file__).resolve().parents[2] / "data"

try:
    import jieba

    jieba.setLogLevel(20)
except ImportError:  # pragma: no cover - 依赖缺失时走正则分词
    jieba = None


SEMANTIC_MODEL = "BAAI/bge-small-zh-v1.5"
DEFAULT_CHUNK_CHARS = 1100
DEFAULT_CHUNK_OVERLAP = 120

_configured_db = os.getenv("DRAGON_KNOWLEDGE_DB", "").strip()
DRAGON_KNOWLEDGE_DB_PATH = Path(_configured_db) if _configured_db else Path(DATA_DIR) / "dragon_knowledge.db"

_EMBEDDING_MODELS: dict[tuple[str, str], object] = {}
_EMBEDDING_LOCK = threading.RLock()

_FIELD_ALIASES: dict[str, tuple[str, ...]] = {
    "stock_code": ("股票代码", "证券代码", "代码", "stock_code", "symbol", "ticker"),
    "stock_name": ("股票名称", "证券名称", "名称", "stock_name", "name"),
    "model_name": ("模型名称", "模型名", "模型", "model_name", "model"),
    "historical_recognition": ("历史辨识度", "辨识度", "historical_recognition", "recognition"),
    "historical_highest_board": ("历史最高板", "最高板", "最高连板", "historical_highest_board"),
    "trigger_conditions": ("模型触发条件", "触发条件", "触发", "trigger_conditions", "trigger"),
    "expectation_points": ("超预期点", "超预期", "expectation_points", "expectation"),
    "success_conditions": ("成功条件", "成功", "success_conditions", "success"),
    "failure_conditions": ("失败条件", "失效条件", "失败", "failure_conditions", "failure"),
    "case_tags": ("案例标签", "标签", "分类", "类型", "case_tags", "tags", "category"),
    "case_title": ("案例标题", "标题", "case_title"),
    "start_date": ("起始日期", "开始日期", "start_date"),
    "end_date": ("结束日期", "截止日期", "end_date"),
    "explicit_attributes": ("明确属性", "历史属性", "属性", "explicit_attributes"),
}

_DISPLAY_FIELDS: dict[str, str] = {
    "stock_code": "股票代码",
    "stock_name": "股票名称",
    "model_name": "模型名称",
    "historical_recognition": "历史辨识度",
    "historical_highest_board": "历史最高板",
    "trigger_conditions": "模型触发条件",
    "expectation_points": "超预期点",
    "success_conditions": "成功条件",
    "failure_conditions": "失败条件",
}

_LABEL_PATTERN_FIELDS = tuple(
    field_name for field_name in _FIELD_ALIASES if field_name != "case_tags"
)
_STOCK_CODE_RE = re.compile(r"(?<!\d)(?:[0368]\d{5}|\d{6})(?!\d)")
_INFER_STOCK_CODE_RE = re.compile(
    r"(?<!\d)(?:(?:000|001|002|003|300|301|600|601|603|605|688|689)\d{3}|(?:4|8|9)\d{5})(?!\d)"
)
_CASE_HEADING_RE = re.compile(
    r"^(?P<name>[\u4e00-\u9fffA-Za-z0-9*·]{2,16})\s+"
    r"(?P<start>(?:\d{4}[-/.年])?\d{1,2}[-/.月]\d{1,2}日?)\s*[-—至~]\s*"
    r"(?P<end>(?:\d{4}[-/.年])?\d{1,2}[-/.月]\d{1,2}日?)\s+"
    r"(?P<boards>\d+(?:\.\d+)?)\s*板?(?=\s|[（(]|$)"
)
_MONTH_HEADING_RE = re.compile(r"^\d{4}年\d{1,2}月$")
_LABELED_STOCK_CASE_RE = re.compile(
    r"^(?:个股|股票名称|证券名称)\s*[：:]\s*(?P<name>[\u4e00-\u9fffA-Za-z0-9*·]{2,16})"
)
_CASE_TITLE_LINE_RE = re.compile(r"^案例标题\s*[：:]\s*(?P<title>.+)$")
_YEAR_RE = re.compile(r"(?<!\d)(20\d{2})(?!\d)")
_CJK_RE = re.compile(r"[\u4e00-\u9fff]{2,}")
_ASCII_WORD_RE = re.compile(r"[A-Za-z][A-Za-z0-9_.-]{1,}")


@dataclass(slots=True)
class ParsedUnit:
    """一个可独立检索的原始记录（例如表格中的一行）。"""

    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


def _now() -> str:
    return datetime.now().isoformat(timespec="seconds")


def _normalize_text(value: Any) -> str:
    text = "" if value is None else str(value)
    if _project_normalize_text is not None:
        return _project_normalize_text(text)
    text = text.replace("\ufeff", "").replace("\u200b", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _content_hash(*values: Any) -> str:
    if _project_content_hash is not None:
        return _project_content_hash(*(str(value) for value in values))
    import hashlib

    payload = "\x1f".join(_normalize_text(value) for value in values)
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def _split_chunks(value: str) -> list[str]:
    if _project_split_chunks is not None:
        return _project_split_chunks(
            value,
            max_chars=DEFAULT_CHUNK_CHARS,
            overlap_chars=DEFAULT_CHUNK_OVERLAP,
        )
    text = _normalize_text(value)
    if not text:
        return []
    result: list[str] = []
    step = max(1, DEFAULT_CHUNK_CHARS - DEFAULT_CHUNK_OVERLAP)
    for start in range(0, len(text), step):
        result.append(text[start : start + DEFAULT_CHUNK_CHARS])
        if start + DEFAULT_CHUNK_CHARS >= len(text):
            break
    return result


def _json_safe(value: Any) -> Any:
    """将 Excel 单元格、Path 等值转成可写入 SQLite JSON 的形态。"""
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    if isinstance(value, Path):
        return str(value)
    if isinstance(value, Mapping):
        return {str(key): _json_safe(item) for key, item in value.items()}
    if isinstance(value, (list, tuple, set)):
        return [_json_safe(item) for item in value]
    return str(value)


def _json_dumps(value: Any) -> str:
    return json.dumps(_json_safe(value), ensure_ascii=False, sort_keys=True)


def _json_loads(value: str | None, fallback: Any) -> Any:
    if not value:
        return fallback
    try:
        return json.loads(value)
    except (TypeError, ValueError):
        return fallback


def _key(value: Any) -> str:
    return re.sub(r"[\s_\-（）()【】\[\]：:]+", "", str(value or "")).lower()


def _text_value(value: Any) -> str:
    if isinstance(value, (list, tuple, set)):
        return "、".join(part for part in (_text_value(item) for item in value) if part)
    if isinstance(value, Mapping):
        return "；".join(
            f"{key}：{_text_value(item)}" for key, item in value.items() if _text_value(item)
        )
    return _normalize_text(value)


def _first_present(mapping: Mapping[str, Any], aliases: Sequence[str]) -> Any:
    normalized = {_key(key): value for key, value in mapping.items()}
    for alias in aliases:
        value = normalized.get(_key(alias))
        if value not in (None, ""):
            return value
    return None


def _normalize_tags(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        values: Iterable[Any] = re.split(r"[，,；;|/、\n]+", value)
    elif isinstance(value, Mapping):
        values = value.values()
    elif isinstance(value, Iterable):
        values = value
    else:
        values = [value]
    seen: set[str] = set()
    result: list[str] = []
    for item in values:
        text = _normalize_text(item)
        if not text or text in seen:
            continue
        seen.add(text)
        result.append(text[:80])
    return result


def _metadata_tags(metadata: Mapping[str, Any]) -> list[str]:
    tags: list[str] = []
    for key_name, value in metadata.items():
        normalized_key = _key(key_name)
        if (
            normalized_key == _key("case_tags")
            or "标签" in str(key_name)
            or normalized_key in {_key("分类"), _key("类型"), _key("category")}
        ):
            tags.extend(_normalize_tags(value))
    return _normalize_tags(tags)


def _extract_label(text: str, aliases: Sequence[str]) -> str:
    for alias in aliases:
        pattern = re.compile(
            rf"(?:^|\n|[；;])\s*{re.escape(alias)}\s*[：:]\s*([^\n；;]{{1,240}})",
            re.IGNORECASE,
        )
        match = pattern.search(text)
        if match:
            return _normalize_text(match.group(1))
    return ""


def _infer_stock_name(text: str, code: str) -> str:
    if not code:
        return ""
    patterns = (
        rf"([\u4e00-\u9fffA-Za-z]{{2,12}})\s*[（(]?\s*{re.escape(code)}",
        rf"{re.escape(code)}\s*[）)]?\s*([\u4e00-\u9fffA-Za-z]{{2,12}})",
    )
    blocked = {"股票代码", "证券代码", "个股代码", "代码", "股票名称", "证券名称"}
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            candidate = match.group(1)
            if candidate not in blocked:
                return candidate
    return ""


def _normalize_stock_code(value: Any) -> str:
    """保留文本代码的前导零，并兼容 Excel 数值单元格。"""
    text = _normalize_text(value)
    if re.fullmatch(r"\d+\.0+", text):
        text = text.split(".", 1)[0]
    match = _STOCK_CODE_RE.search(text)
    if match:
        return match.group(0)
    if re.fullmatch(r"\d{1,6}", text):
        return text.zfill(6)
    return text


def _normalize_case_date(value: Any, *, year_hint: str = "") -> str:
    text = _normalize_text(value).replace("年", "-").replace("月", "-").replace("日", "")
    text = text.replace("/", "-").replace(".", "-")
    parts = [part for part in text.split("-") if part]
    if len(parts) == 2 and year_hint:
        parts.insert(0, year_hint)
    if len(parts) != 3 or not all(part.isdigit() for part in parts):
        return _normalize_text(value)
    year, month, day = map(int, parts)
    if not (2000 <= year <= 2099 and 1 <= month <= 12 and 1 <= day <= 31):
        return _normalize_text(value)
    return f"{year:04d}-{month:02d}-{day:02d}"


def _case_date_range(metadata: Mapping[str, Any]) -> str:
    start = _text_value(metadata.get("start_date", ""))
    end = _text_value(metadata.get("end_date", ""))
    if start and end:
        return start if start == end else f"{start}—{end}"
    return start or end


def _adjust_cross_year_range(start: str, end: str) -> tuple[str, str]:
    """“12/25-1/03”按跨年处理，避免生成同年倒序日期。"""

    if not (re.fullmatch(r"20\d{2}-\d{2}-\d{2}", start) and re.fullmatch(r"20\d{2}-\d{2}-\d{2}", end)):
        return start, end
    start_year, start_month = int(start[:4]), int(start[5:7])
    end_year, end_month = int(end[:4]), int(end[5:7])
    if start_year == end_year and start_month > end_month:
        start = f"{start_year - 1:04d}{start[4:]}"
    return start, end


def _mentioned_stocks(text: str, owner_name: str = "") -> list[str]:
    values: list[str] = []
    for code in _INFER_STOCK_CODE_RE.findall(text):
        if code not in values:
            values.append(code)
    # “正文提及股票”只作审计线索，不据此把元数据归给其他股票。
    for match in re.finditer(r"([\u4e00-\u9fffA-Za-z*·]{2,12})(?=[（(]?\d{6}[）)]?)", text):
        name = match.group(1).strip()
        if name and name != owner_name and name not in values:
            values.append(name)
    return values


def _extract_metadata(
    text: str,
    supplied: Mapping[str, Any] | None = None,
    *,
    infer_unlabeled_stock: bool = True,
) -> dict[str, Any]:
    """兼容用户表头和自然文本中的主要历史模型字段。"""
    raw: dict[str, Any] = {
        str(key): _json_safe(value) for key, value in (supplied or {}).items() if value not in (None, "")
    }
    normalized_text = _normalize_text(text)
    for field_name, aliases in _FIELD_ALIASES.items():
        value = _first_present(raw, aliases)
        if value in (None, ""):
            value = _extract_label(normalized_text, aliases)
        if value not in (None, ""):
            raw[field_name] = _json_safe(value)

    code = _normalize_stock_code(raw.get("stock_code", ""))
    if not code and infer_unlabeled_stock:
        code_match = _INFER_STOCK_CODE_RE.search(normalized_text)
        code = code_match.group(0) if code_match else ""
    if code:
        raw["stock_code"] = code

    name = _text_value(raw.get("stock_name", ""))
    if not name:
        name = _infer_stock_name(normalized_text, code)
    if name:
        raw["stock_name"] = name

    raw_tags = _normalize_tags(raw.get("case_tags")) + _metadata_tags(raw)
    if raw_tags:
        raw["case_tags"] = _normalize_tags(raw_tags)
    return raw


def _merge_metadata(*values: Mapping[str, Any]) -> dict[str, Any]:
    merged: dict[str, Any] = {}
    for value in values:
        for key_name, item in value.items():
            if item not in (None, ""):
                merged[str(key_name)] = _json_safe(item)
    return _extract_metadata("", merged)


def _tokens(value: str, *, maximum: int = 160) -> list[str]:
    """中文关键词、股票代码和字符二元组的轻量本地召回特征。"""
    text = _normalize_text(value).lower()
    if not text:
        return []
    values: list[str] = []
    if jieba is not None:
        try:
            values.extend(
                word.strip().lower()
                for word in jieba.lcut(text)
                if word.strip() and (len(word.strip()) > 1 or word.strip().isdigit())
            )
        except Exception:  # pragma: no cover - 词典初始化异常时仍可检索
            pass
    values.extend(match.group(0).lower() for match in _ASCII_WORD_RE.finditer(text))
    values.extend(match.group(0) for match in _STOCK_CODE_RE.finditer(text))
    for phrase in _CJK_RE.findall(text):
        values.append(phrase)
        values.extend(phrase[index : index + 2] for index in range(max(0, len(phrase) - 1)))
    seen: set[str] = set()
    result: list[str] = []
    for item in values:
        item = item.strip()
        if not item or item in seen or len(item) > 80:
            continue
        seen.add(item)
        result.append(item)
        if len(result) >= maximum:
            break
    return result


def _counter(tokens: Iterable[str]) -> Counter[str]:
    return Counter(token for token in tokens if token)


def _cosine(left: Counter[str], right: Counter[str]) -> float:
    if not left or not right:
        return 0.0
    numerator = sum(left[key] * right[key] for key in set(left) & set(right))
    denominator = math.sqrt(sum(item * item for item in left.values())) * math.sqrt(
        sum(item * item for item in right.values())
    )
    return numerator / denominator if denominator else 0.0


def _normalise_vector(values: Iterable[Any]) -> array:
    vector = array("f", (float(item) for item in values))
    length = math.sqrt(sum(item * item for item in vector))
    if length:
        for index, item in enumerate(vector):
            vector[index] = item / length
    return vector


def _vector_from_blob(blob: bytes, dimensions: int) -> array | None:
    vector = array("f")
    try:
        vector.frombytes(blob)
    except (EOFError, ValueError):
        return None
    return vector if len(vector) == dimensions else None


def _dot(left: array, right: array) -> float:
    if len(left) != len(right) or not left:
        return 0.0
    return sum(a * b for a, b in zip(left, right))


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "utf-16"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _normalise_header(value: Any, index: int) -> str:
    text = _normalize_text(value)
    return text or f"字段{index + 1}"


def _tabular_units(rows: Sequence[Sequence[Any]], *, scope: str) -> list[ParsedUnit]:
    materialized = [list(row) for row in rows if any(item not in (None, "") for item in row)]
    if not materialized:
        return []
    header_index = 0
    for index, row in enumerate(materialized[:10]):
        if sum(item not in (None, "") for item in row) >= 2:
            header_index = index
            break
    headers = [_normalise_header(value, index) for index, value in enumerate(materialized[header_index])]
    prefix_lines = [
        _normalize_text(" | ".join(_text_value(value) for value in row if value not in (None, "")))
        for row in materialized[:header_index]
    ]
    prefix_lines = [value for value in prefix_lines if value]
    units: list[ParsedUnit] = []
    for row_number, row in enumerate(materialized[header_index + 1 :], header_index + 2):
        mapping = {
            headers[index]: _json_safe(value)
            for index, value in enumerate(row)
            if index < len(headers) and value not in (None, "")
        }
        if not mapping:
            continue
        parts = [f"{key_name}：{_text_value(value)}" for key_name, value in mapping.items()]
        lead = f"{scope}；第{row_number}行"
        if prefix_lines:
            lead += f"；{'；'.join(prefix_lines)}"
        units.append(ParsedUnit(text=f"{lead}\n" + "\n".join(parts), metadata=mapping))
    if not units:
        fallback = "\n".join(
            " | ".join(_text_value(value) for value in row if value not in (None, ""))
            for row in materialized
        )
        if fallback:
            units.append(ParsedUnit(text=f"{scope}\n{fallback}"))
    return units


def _walk_docx_tables(tables: Iterable[Any]) -> Iterable[Any]:
    for table in tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _walk_docx_tables(cell.tables)


def _parse_docx(content: bytes) -> list[ParsedUnit]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - 依赖不完整的运行环境
        raise RuntimeError("解析 DOCX 需要 python-docx") from exc
    document = Document(io.BytesIO(content))
    units: list[ParsedUnit] = []
    current_heading = ""
    current_case_heading = ""
    current_case_metadata: dict[str, Any] = {}
    current_year = ""
    case_index = 0
    for paragraph in document.paragraphs:
        text = _normalize_text(paragraph.text)
        if not text:
            continue
        year_match = _YEAR_RE.search(text)
        if year_match:
            current_year = year_match.group(1)
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading") or style_name.startswith("标题"):
            current_heading = text
            year_match = _YEAR_RE.search(text)
            if year_match:
                current_year = year_match.group(1)
            current_case_heading = ""
            current_case_metadata = {}
            continue
        case_match = _CASE_HEADING_RE.match(text)
        labeled_stock = _LABELED_STOCK_CASE_RE.match(text)
        if labeled_stock:
            case_index += 1
            current_case_heading = text
            current_case_metadata = {
                "stock_name": labeled_stock.group("name").strip(),
                "case_title": text,
                "case_key": f"docx-case-{case_index}",
                "chunk_role": "案例标题",
                "extraction_status": "已提取",
            }
            prefix = f"章节：{current_heading}\n" if current_heading else ""
            units.append(ParsedUnit(text=prefix + text, metadata=dict(current_case_metadata)))
            continue
        title_match = _CASE_TITLE_LINE_RE.match(text)
        if title_match and current_case_metadata:
            current_case_heading = title_match.group("title").strip()
            current_case_metadata["case_title"] = current_case_heading
            title_metadata = dict(current_case_metadata)
            title_metadata["chunk_role"] = "案例标题"
            prefix = f"章节：{current_heading}\n" if current_heading else ""
            units.append(ParsedUnit(text=prefix + text, metadata=title_metadata))
            continue
        if case_match:
            case_index += 1
            current_case_heading = text
            year_match = _YEAR_RE.search(text)
            year_hint = year_match.group(1) if year_match else current_year
            start_date = _normalize_case_date(case_match.group("start"), year_hint=year_hint)
            end_date = _normalize_case_date(case_match.group("end"), year_hint=year_hint)
            start_date, end_date = _adjust_cross_year_range(start_date, end_date)
            current_case_metadata = {
                "stock_name": case_match.group("name").strip(),
                "historical_highest_board": case_match.group("boards"),
                "case_title": text,
                "start_date": start_date,
                "end_date": end_date,
                "case_date": start_date if start_date == end_date else f"{start_date}—{end_date}",
                "case_key": f"docx-case-{case_index}",
                "chunk_role": "案例标题",
                "extraction_status": "已提取",
            }
            prefix = f"章节：{current_heading}\n" if current_heading else ""
            units.append(ParsedUnit(text=prefix + text, metadata=dict(current_case_metadata)))
            continue
        if _MONTH_HEADING_RE.fullmatch(text):
            year_match = _YEAR_RE.search(text)
            if year_match:
                current_year = year_match.group(1)
            current_case_heading = ""
            current_case_metadata = {}
            continue
        prefix = f"章节：{current_heading}\n" if current_heading else ""
        if current_case_metadata:
            prefix += f"个股：{current_case_metadata['stock_name']}\n案例标题：{current_case_heading}\n"
        paragraph_metadata = dict(current_case_metadata)
        if paragraph_metadata:
            paragraph_metadata["chunk_role"] = "正文"
            paragraph_metadata["mentioned_stocks"] = _mentioned_stocks(
                text, paragraph_metadata.get("stock_name", "")
            )
        units.append(ParsedUnit(text=prefix + text, metadata=paragraph_metadata))
    seen_tables: set[int] = set()
    for table_index, table in enumerate(_walk_docx_tables(document.tables), 1):
        table_key = id(table._tbl)
        if table_key in seen_tables:
            continue
        seen_tables.add(table_key)
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        units.extend(_tabular_units(rows, scope=f"DOCX 表格{table_index}"))
    return units


def _parse_xlsx(content: bytes) -> list[ParsedUnit]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover - 依赖不完整的运行环境
        raise RuntimeError("解析 XLSX 需要 openpyxl") from exc
    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    units: list[ParsedUnit] = []
    try:
        for sheet in workbook.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            units.extend(_tabular_units(rows, scope=f"工作表：{sheet.title}"))
    finally:
        workbook.close()
    return units


def _parse_csv(content: bytes) -> list[ParsedUnit]:
    text = _decode_text(content)
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
    except csv.Error:
        dialect = csv.excel
    rows = list(csv.reader(io.StringIO(text), dialect))
    return _tabular_units(rows, scope="CSV")


def _mapping_unit(mapping: Mapping[str, Any], *, scope: str) -> ParsedUnit:
    metadata = {str(key): _json_safe(value) for key, value in mapping.items()}
    parts = [
        f"{key_name}：{_text_value(value)}"
        for key_name, value in metadata.items()
        if _text_value(value)
    ]
    return ParsedUnit(text=f"{scope}\n" + "\n".join(parts), metadata=metadata)


def _parse_json(content: bytes) -> list[ParsedUnit]:
    try:
        data = json.loads(_decode_text(content))
    except json.JSONDecodeError as exc:
        raise ValueError(f"JSON 格式无效：{exc.msg}") from exc
    if isinstance(data, list):
        return [
            _mapping_unit(item, scope=f"JSON 记录{index}")
            if isinstance(item, Mapping)
            else ParsedUnit(text=f"JSON 记录{index}\n{_text_value(item)}")
            for index, item in enumerate(data, 1)
            if _text_value(item)
        ]
    if isinstance(data, Mapping):
        for key_name in ("documents", "records", "items", "data", "cases", "models"):
            candidate = data.get(key_name)
            if isinstance(candidate, list):
                shared = {key: value for key, value in data.items() if key != key_name and not isinstance(value, list)}
                units: list[ParsedUnit] = []
                for index, item in enumerate(candidate, 1):
                    if isinstance(item, Mapping):
                        units.append(_mapping_unit({**shared, **item}, scope=f"JSON {key_name} {index}"))
                    elif _text_value(item):
                        units.append(ParsedUnit(text=f"JSON {key_name} {index}\n{_text_value(item)}", metadata=shared))
                return units
        return [_mapping_unit(data, scope="JSON")]
    return [ParsedUnit(text=f"JSON\n{_text_value(data)}")]


def _parse_markdown_or_text(content: bytes) -> list[ParsedUnit]:
    text = _normalize_text(_decode_text(content))
    if not text:
        return []
    metadata: dict[str, Any] = {}
    if text.startswith("---\n"):
        closing = text.find("\n---", 4)
        if closing > 0:
            for line in text[4:closing].splitlines():
                if ":" in line or "：" in line:
                    key_name, value = re.split(r"[:：]", line, maxsplit=1)
                    if key_name.strip() and value.strip():
                        metadata[key_name.strip()] = value.strip()
            text = _normalize_text(text[closing + 4 :])
    return [ParsedUnit(text=text, metadata=metadata)]


def _parse_file_content(filename: str, content: bytes) -> tuple[str, list[ParsedUnit]]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return "docx", _parse_docx(content)
    if suffix == ".xlsx":
        return "xlsx", _parse_xlsx(content)
    if suffix == ".csv":
        return "csv", _parse_csv(content)
    if suffix == ".json":
        return "json", _parse_json(content)
    if suffix in {".md", ".markdown"}:
        return "markdown", _parse_markdown_or_text(content)
    if suffix in {".txt", ".text"}:
        return "text", _parse_markdown_or_text(content)
    raise ValueError("仅支持 DOCX、XLSX、CSV、JSON、Markdown 和 TXT 文件")


class DragonKnowledgeStore:
    """首板布局的独立 RAG 存储。

    可直接由路由层调用的主要方法：

    * :meth:`import_file` / :meth:`import_bytes` / :meth:`import_text`
    * :meth:`list_documents`、:meth:`list_tags`、:meth:`update_document_tags`
    * :meth:`search`

    ``review_knowledge.db`` 被显式拦截，避免误把首板资料写入旧复盘 RAG。
    """

    def __init__(
        self,
        path: str | Path | None = None,
        *,
        semantic_model: str = SEMANTIC_MODEL,
        semantic_cache_dir: str | Path | None = None,
    ) -> None:
        candidate = Path(path) if path is not None else DRAGON_KNOWLEDGE_DB_PATH
        resolved = candidate.expanduser().resolve()
        if resolved.name.casefold() != "dragon_knowledge.db":
            raise ValueError("首板布局知识库必须使用独立的 dragon_knowledge.db")
        self.path = resolved
        self.path.parent.mkdir(parents=True, exist_ok=True)
        self.database_path = str(self.path)
        self.semantic_model = semantic_model
        shared_cache = Path(DATA_DIR) / "fastembed_cache"
        self.semantic_cache_dir = (
            Path(semantic_cache_dir).expanduser().resolve()
            if semantic_cache_dir is not None
            else (shared_cache if shared_cache.exists() else self.path.parent / "fastembed_cache")
        )
        self._lock = threading.RLock()
        self.connection = sqlite3.connect(self.path, check_same_thread=False)
        self.connection.row_factory = sqlite3.Row
        self.connection.execute("PRAGMA foreign_keys=ON")
        self.connection.execute("PRAGMA busy_timeout=5000")
        self.fts_available = False
        self._semantic_error = ""
        self._create_schema()

    def __enter__(self) -> "DragonKnowledgeStore":
        return self

    def __exit__(self, *_args: Any) -> None:
        self.close()

    def close(self) -> None:
        with self._lock:
            self.connection.close()

    def _create_schema(self) -> None:
        with self._lock:
            self.connection.executescript(
                """
                PRAGMA journal_mode=WAL;
                CREATE TABLE IF NOT EXISTS documents (
                    id TEXT PRIMARY KEY,
                    source_path TEXT NOT NULL UNIQUE,
                    title TEXT NOT NULL,
                    source_type TEXT NOT NULL,
                    content_hash TEXT NOT NULL,
                    metadata_json TEXT NOT NULL DEFAULT '{}',
                    tags_json TEXT NOT NULL DEFAULT '[]',
                    raw_text TEXT NOT NULL DEFAULT '',
                    created_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL
                );
                CREATE TABLE IF NOT EXISTS chunks (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    document_id TEXT NOT NULL,
                    chunk_index INTEGER NOT NULL,
                    content TEXT NOT NULL,
                    content_hash TEXT NOT NULL,
                    search_text TEXT NOT NULL,
                    metadata_json TEXT NOT NULL DEFAULT '{}',
                    tags_json TEXT NOT NULL DEFAULT '[]',
                    stock_code TEXT NOT NULL DEFAULT '',
                    stock_name TEXT NOT NULL DEFAULT '',
                    model_name TEXT NOT NULL DEFAULT '',
                    parent_case_id TEXT NOT NULL DEFAULT '',
                    chunk_role TEXT NOT NULL DEFAULT '正文',
                    mentioned_stocks_json TEXT NOT NULL DEFAULT '[]',
                    mentioned_attributes_json TEXT NOT NULL DEFAULT '[]',
                    mentioned_highest_board TEXT NOT NULL DEFAULT '',
                    created_at TEXT NOT NULL,
                    FOREIGN KEY(document_id) REFERENCES documents(id) ON DELETE CASCADE,
                    UNIQUE(document_id, chunk_index)
                );
                CREATE INDEX IF NOT EXISTS idx_chunks_document ON chunks(document_id, chunk_index);
                CREATE INDEX IF NOT EXISTS idx_chunks_stock_code ON chunks(stock_code);
                CREATE INDEX IF NOT EXISTS idx_chunks_stock_name ON chunks(stock_name);
                CREATE TABLE IF NOT EXISTS cases (
                    id TEXT PRIMARY KEY,
                    document_id TEXT NOT NULL,
                    case_index INTEGER NOT NULL,
                    stock_code TEXT NOT NULL DEFAULT '',
                    stock_name TEXT NOT NULL DEFAULT '',
                    case_title TEXT NOT NULL DEFAULT '',
                    historical_highest_board TEXT NOT NULL DEFAULT '',
                    model_name TEXT NOT NULL DEFAULT '',
                    explicit_attributes_json TEXT NOT NULL DEFAULT '[]',
                    start_date TEXT NOT NULL DEFAULT '',
                    end_date TEXT NOT NULL DEFAULT '',
                    case_text TEXT NOT NULL DEFAULT '',
                    extraction_status TEXT NOT NULL DEFAULT '已提取',
                    manual_correction_status TEXT NOT NULL DEFAULT '未修正',
                    metadata_json TEXT NOT NULL DEFAULT '{}',
                    created_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL,
                    FOREIGN KEY(document_id) REFERENCES documents(id) ON DELETE CASCADE,
                    UNIQUE(document_id, case_index)
                );
                CREATE INDEX IF NOT EXISTS idx_cases_stock_code ON cases(stock_code);
                CREATE INDEX IF NOT EXISTS idx_cases_stock_name ON cases(stock_name);
                CREATE TABLE IF NOT EXISTS embeddings (
                    chunk_id INTEGER NOT NULL,
                    model TEXT NOT NULL,
                    content_hash TEXT NOT NULL,
                    dimensions INTEGER NOT NULL,
                    vector BLOB NOT NULL,
                    updated_at TEXT NOT NULL,
                    PRIMARY KEY(chunk_id, model),
                    FOREIGN KEY(chunk_id) REFERENCES chunks(id) ON DELETE CASCADE
                );
                CREATE TABLE IF NOT EXISTS knowledge_meta (
                    key TEXT PRIMARY KEY,
                    value TEXT NOT NULL
                );
                """
            )
            # 兼容已经存在的 dragon_knowledge.db；只补列，不触碰原始资料文件。
            existing_chunk_columns = {
                row["name"] for row in self.connection.execute("PRAGMA table_info(chunks)").fetchall()
            }
            migrations = {
                "parent_case_id": "TEXT NOT NULL DEFAULT ''",
                "chunk_role": "TEXT NOT NULL DEFAULT '正文'",
                "mentioned_stocks_json": "TEXT NOT NULL DEFAULT '[]'",
                "mentioned_attributes_json": "TEXT NOT NULL DEFAULT '[]'",
                "mentioned_highest_board": "TEXT NOT NULL DEFAULT ''",
            }
            for column, definition in migrations.items():
                if column not in existing_chunk_columns:
                    self.connection.execute(f"ALTER TABLE chunks ADD COLUMN {column} {definition}")
            self.connection.execute(
                "CREATE INDEX IF NOT EXISTS idx_chunks_parent_case ON chunks(parent_case_id, chunk_index)"
            )
            try:
                self.connection.execute(
                    "CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts USING fts5(chunk_id UNINDEXED, indexed_text)"
                )
                self.fts_available = True
            except sqlite3.OperationalError:
                self.fts_available = False
            self.connection.commit()

    @staticmethod
    def _source_title(source_path: str, supplied: str | None, metadata: Mapping[str, Any]) -> str:
        if supplied and _normalize_text(supplied):
            return _normalize_text(supplied)
        model_name = _text_value(metadata.get("model_name", ""))
        if model_name:
            return model_name
        path = Path(source_path.replace("upload://", ""))
        return path.stem or "未命名历史资料"

    @staticmethod
    def _result_document(row: sqlite3.Row | Mapping[str, Any], *, chunks: int = 0, embeddings: int = 0) -> dict[str, Any]:
        data = dict(row)
        tags = _normalize_tags(_json_loads(data.pop("tags_json", "[]"), []))
        metadata = _json_loads(data.pop("metadata_json", "{}"), {})
        result = {
            "id": data["id"],
            "document_id": data["id"],
            "source_id": data["id"],
            "title": data["title"],
            "source_path": data["source_path"],
            "source_type": data["source_type"],
            "content_hash": data["content_hash"],
            "tags": tags,
            "metadata": metadata if isinstance(metadata, dict) else {},
            "created_at": data["created_at"],
            "updated_at": data["updated_at"],
            "chunks": int(chunks),
            "embeddings": int(embeddings),
        }
        for field_name in _DISPLAY_FIELDS:
            if field_name in result["metadata"]:
                result[field_name] = result["metadata"][field_name]
        return result

    def _fts_text(
        self,
        *,
        title: str,
        content: str,
        metadata: Mapping[str, Any],
        tags: Sequence[str],
    ) -> str:
        metadata_text = "\n".join(
            f"{key_name}：{_text_value(value)}" for key_name, value in metadata.items() if _text_value(value)
        )
        raw = "\n".join((title, content, metadata_text, " ".join(tags)))
        return " ".join(_tokens(raw, maximum=500))

    def _prepared_chunks(
        self,
        units: Sequence[ParsedUnit],
        *,
        title: str,
        document_metadata: Mapping[str, Any],
        document_tags: Sequence[str],
    ) -> list[dict[str, Any]]:
        prepared: list[dict[str, Any]] = []
        for unit_index, unit in enumerate(units):
            text = _normalize_text(unit.text)
            if not text:
                continue
            inferred = _extract_metadata(
                text,
                unit.metadata,
                infer_unlabeled_stock=(
                    len(text) <= 2_000
                    and len(set(_INFER_STOCK_CODE_RE.findall(text))) <= 1
                ),
            )
            metadata = _merge_metadata(document_metadata, inferred)
            case_key = _text_value(metadata.get("case_key", "")) or f"unit-{unit_index}"
            metadata["case_key"] = case_key
            if not metadata.get("case_title"):
                metadata["case_title"] = _text_value(metadata.get("stock_name", "")) or title
            if not metadata.get("case_date"):
                metadata["case_date"] = _case_date_range(metadata)
            chunk_tags = _normalize_tags([*document_tags, *_metadata_tags(metadata)])
            for fragment in _split_chunks(text):
                fragment = _normalize_text(fragment)
                if not fragment:
                    continue
                prepared.append(
                    {
                        "content": fragment,
                        "content_hash": _content_hash(fragment),
                        "metadata": metadata,
                        "tags": chunk_tags,
                        "stock_code": _text_value(metadata.get("stock_code", "")),
                        "stock_name": _text_value(metadata.get("stock_name", "")),
                        "model_name": _text_value(metadata.get("model_name", "")),
                        "case_key": case_key,
                        "chunk_role": _text_value(metadata.get("chunk_role", "正文")) or "正文",
                        "mentioned_stocks": _normalize_tags(metadata.get("mentioned_stocks", [])),
                        "mentioned_attributes": _normalize_tags(metadata.get("mentioned_attributes", [])),
                        "mentioned_highest_board": _text_value(
                            metadata.get("mentioned_highest_board", "")
                        ),
                        "search_text": self._fts_text(
                            title=title,
                            content=fragment,
                            metadata=metadata,
                            tags=chunk_tags,
                        ),
                    }
                )
        return prepared

    def import_file(
        self,
        path: str | Path,
        *,
        title: str | None = None,
        tags: Sequence[str] | str | None = None,
        metadata: Mapping[str, Any] | None = None,
        build_semantic: bool = True,
    ) -> dict[str, Any]:
        """解析本地资料并覆盖更新同一路径的来源记录。"""
        source_path = Path(path).expanduser().resolve()
        if not source_path.is_file():
            raise FileNotFoundError(f"未找到历史模型资料：{source_path}")
        source_type, units = _parse_file_content(source_path.name, source_path.read_bytes())
        return self.import_text(
            "\n\n".join(unit.text for unit in units),
            title=title,
            source_path=str(source_path),
            source_type=source_type,
            tags=tags,
            metadata=metadata,
            units=units,
            build_semantic=build_semantic,
        )

    def import_bytes(
        self,
        filename: str,
        content: bytes,
        *,
        title: str | None = None,
        tags: Sequence[str] | str | None = None,
        metadata: Mapping[str, Any] | None = None,
        source_id: str | None = None,
        build_semantic: bool = True,
    ) -> dict[str, Any]:
        """供 API 上传层调用；不需要将临时上传文件写入旧知识库目录。"""
        safe_name = Path(filename or "历史模型资料.txt").name
        source_type, units = _parse_file_content(safe_name, content)
        identity = source_id or str(uuid4())
        return self.import_text(
            "\n\n".join(unit.text for unit in units),
            title=title,
            source_path=f"upload://{identity}/{safe_name}",
            source_type=source_type,
            tags=tags,
            metadata=metadata,
            units=units,
            build_semantic=build_semantic,
        )

    def import_text(
        self,
        text: str,
        *,
        title: str | None = None,
        source_path: str | Path | None = None,
        source_type: str = "text",
        tags: Sequence[str] | str | None = None,
        metadata: Mapping[str, Any] | None = None,
        units: Sequence[ParsedUnit] | None = None,
        build_semantic: bool = True,
    ) -> dict[str, Any]:
        """导入已解析文本。``source_path`` 相同的资料会原位更新。"""
        content = _normalize_text(text)
        if not content:
            raise ValueError("历史模型资料没有可索引的文本")
        source = str(source_path) if source_path is not None else f"inline://{uuid4()}"
        source_units = list(units) if units is not None else [ParsedUnit(content)]
        # 文档级元数据只接受明确提供的字段；整篇正文可能包含大量股票与六位数，
        # 不能把第一个偶然命中扩散到所有切片。具体个股字段在解析单元级提取。
        document_metadata = _extract_metadata("", metadata, infer_unlabeled_stock=False)
        supplied_tags = _normalize_tags(tags)
        previous_tags: list[str] = []
        if tags is None:
            with self._lock:
                previous = self.connection.execute(
                    "SELECT tags_json FROM documents WHERE source_path = ?", (source,)
                ).fetchone()
            if previous is not None:
                previous_tags = _normalize_tags(_json_loads(previous["tags_json"], []))
        document_tags = _normalize_tags([*supplied_tags, *previous_tags, *_metadata_tags(document_metadata)])
        resolved_title = self._source_title(source, title, document_metadata)
        prepared = self._prepared_chunks(
            source_units,
            title=resolved_title,
            document_metadata=document_metadata,
            document_tags=document_tags,
        )
        if not prepared:
            raise ValueError("历史模型资料没有可索引的文本")
        now = _now()
        digest = _content_hash(content, _json_dumps(document_metadata))
        with self._lock, self.connection:
            existing = self.connection.execute(
                "SELECT id, created_at FROM documents WHERE source_path = ?", (source,)
            ).fetchone()
            document_id = existing["id"] if existing else str(uuid4())
            created_at = existing["created_at"] if existing else now
            self.connection.execute(
                """
                INSERT INTO documents(
                    id, source_path, title, source_type, content_hash, metadata_json,
                    tags_json, raw_text, created_at, updated_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(source_path) DO UPDATE SET
                    title=excluded.title, source_type=excluded.source_type,
                    content_hash=excluded.content_hash, metadata_json=excluded.metadata_json,
                    tags_json=excluded.tags_json, raw_text=excluded.raw_text,
                    updated_at=excluded.updated_at
                """,
                (
                    document_id,
                    source,
                    resolved_title,
                    source_type,
                    digest,
                    _json_dumps(document_metadata),
                    _json_dumps(document_tags),
                    content,
                    created_at,
                    now,
                ),
            )
            case_groups: dict[str, list[dict[str, Any]]] = {}
            for item in prepared:
                case_groups.setdefault(item["case_key"], []).append(item)
            existing_cases = {
                int(row["case_index"]): row
                for row in self.connection.execute(
                    "SELECT id, case_index, created_at FROM cases WHERE document_id = ?",
                    (document_id,),
                ).fetchall()
            }
            case_rows: list[tuple[Any, ...]] = []
            for case_index, (_case_key, items) in enumerate(case_groups.items()):
                case_metadata = dict(items[0]["metadata"])
                existing_case = existing_cases.get(case_index)
                case_id = (
                    existing_case["id"]
                    if existing_case is not None
                    else f"case_{_content_hash(document_id, _case_key)[:24]}"
                )
                case_created_at = existing_case["created_at"] if existing_case is not None else now
                case_text = "\n".join(dict.fromkeys(item["content"] for item in items))
                explicit_attributes = _normalize_tags(
                    case_metadata.get("explicit_attributes", case_metadata.get("case_tags", []))
                )
                case_rows.append(
                    (
                        case_id,
                        document_id,
                        case_index,
                        _text_value(case_metadata.get("stock_code", "")),
                        _text_value(case_metadata.get("stock_name", "")),
                        _text_value(case_metadata.get("case_title", "")),
                        _text_value(case_metadata.get("historical_highest_board", "")),
                        _text_value(case_metadata.get("model_name", "")),
                        _json_dumps(explicit_attributes),
                        _text_value(case_metadata.get("start_date", "")),
                        _text_value(case_metadata.get("end_date", "")),
                        case_text,
                        _text_value(case_metadata.get("extraction_status", "已提取")) or "已提取",
                        _text_value(case_metadata.get("manual_correction_status", "未修正")) or "未修正",
                        _json_dumps(case_metadata),
                        case_created_at,
                        now,
                    )
                )
                for item in items:
                    item["parent_case_id"] = case_id
            self.connection.executemany(
                """
                INSERT INTO cases(
                    id, document_id, case_index, stock_code, stock_name, case_title,
                    historical_highest_board, model_name, explicit_attributes_json,
                    start_date, end_date, case_text, extraction_status,
                    manual_correction_status, metadata_json, created_at, updated_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(document_id, case_index) DO UPDATE SET
                    stock_code=excluded.stock_code, stock_name=excluded.stock_name,
                    case_title=excluded.case_title,
                    historical_highest_board=excluded.historical_highest_board,
                    model_name=excluded.model_name,
                    explicit_attributes_json=excluded.explicit_attributes_json,
                    start_date=excluded.start_date, end_date=excluded.end_date,
                    case_text=excluded.case_text,
                    extraction_status=excluded.extraction_status,
                    manual_correction_status=excluded.manual_correction_status,
                    metadata_json=excluded.metadata_json, updated_at=excluded.updated_at
                """,
                case_rows,
            )
            self.connection.execute(
                "DELETE FROM cases WHERE document_id = ? AND case_index >= ?",
                (document_id, len(case_rows)),
            )
            self.connection.executemany(
                """
                INSERT INTO chunks(
                    document_id, chunk_index, content, content_hash, search_text,
                    metadata_json, tags_json, stock_code, stock_name, model_name,
                    parent_case_id, chunk_role, mentioned_stocks_json,
                    mentioned_attributes_json, mentioned_highest_board, created_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(document_id, chunk_index) DO UPDATE SET
                    content=excluded.content,
                    content_hash=excluded.content_hash,
                    search_text=excluded.search_text,
                    metadata_json=excluded.metadata_json,
                    tags_json=excluded.tags_json,
                    stock_code=excluded.stock_code,
                    stock_name=excluded.stock_name,
                    model_name=excluded.model_name,
                    parent_case_id=excluded.parent_case_id,
                    chunk_role=excluded.chunk_role,
                    mentioned_stocks_json=excluded.mentioned_stocks_json,
                    mentioned_attributes_json=excluded.mentioned_attributes_json,
                    mentioned_highest_board=excluded.mentioned_highest_board
                """,
                [
                    (
                        document_id,
                        index,
                        item["content"],
                        item["content_hash"],
                        item["search_text"],
                        _json_dumps(item["metadata"]),
                        _json_dumps(item["tags"]),
                        item["stock_code"],
                        item["stock_name"],
                        item["model_name"],
                        item["parent_case_id"],
                        item["chunk_role"],
                        _json_dumps(item["mentioned_stocks"]),
                        _json_dumps(item["mentioned_attributes"]),
                        item["mentioned_highest_board"],
                        now,
                    )
                    for index, item in enumerate(prepared)
                ],
            )
            # 保留相同位置切片的 id，使未变化内容可继续命中向量缓存；
            # 文档缩短时只删除已经不存在的尾部切片及其级联向量。
            self.connection.execute(
                "DELETE FROM chunks WHERE document_id = ? AND chunk_index >= ?",
                (document_id, len(prepared)),
            )
            self._rebuild_fts_unlocked()
        semantic_result = (
            self.build_semantic_index(document_id=document_id) if build_semantic else self.semantic_status()
        )
        result = self.get_document(document_id)
        assert result is not None  # 已在同一事务写入
        result.update(imported=True, chunks=len(prepared), semantic=semantic_result)
        return result

    # 兼容路由层和后续业务代码的更直观命名。
    ingest_file = import_file
    ingest_text = import_text
    add_document = import_file

    def _rebuild_fts_unlocked(self) -> None:
        if not self.fts_available:
            return
        self.connection.execute("DELETE FROM chunks_fts")
        self.connection.execute(
            "INSERT INTO chunks_fts(chunk_id, indexed_text) SELECT id, search_text FROM chunks"
        )

    def rebuild_fts(self) -> None:
        with self._lock, self.connection:
            self._rebuild_fts_unlocked()

    def _has_local_semantic_model(self) -> bool:
        if not self.semantic_cache_dir.exists():
            return False
        slug = self.semantic_model.rsplit("/", 1)[-1].lower()
        return any(
            slug in str(path).lower() and path.name.lower().endswith(".onnx")
            for path in self.semantic_cache_dir.rglob("*.onnx")
        )

    def _embedding_model(self) -> object | None:
        """仅在已有本地模型文件时加载，避免上传/检索触发网络下载。"""
        if not self._has_local_semantic_model():
            self._semantic_error = "未发现本机 FastEmbed 模型缓存"
            return None
        cache_key = (self.semantic_model, str(self.semantic_cache_dir))
        with _EMBEDDING_LOCK:
            if cache_key in _EMBEDDING_MODELS:
                return _EMBEDDING_MODELS[cache_key]
            try:
                from fastembed import TextEmbedding

                model = TextEmbedding(
                    model_name=self.semantic_model,
                    cache_dir=str(self.semantic_cache_dir),
                    lazy_load=True,
                )
                _EMBEDDING_MODELS[cache_key] = model
                self._semantic_error = ""
                return model
            except Exception as exc:  # 模型、ONNX 或运行时不完整时保持关键词检索可用
                self._semantic_error = f"FastEmbed 未启用：{exc}"
                return None

    def _embed(self, texts: Sequence[str]) -> list[array] | None:
        if not texts:
            return []
        model = self._embedding_model()
        if model is None:
            return None
        try:
            generated = model.embed(list(texts), batch_size=32)  # type: ignore[attr-defined]
            vectors = [_normalise_vector(vector) for vector in generated]
        except Exception as exc:  # 不因本地模型不可用中断上传、FTS 或精确检索
            self._semantic_error = f"FastEmbed 未启用：{exc}"
            return None
        if len(vectors) != len(texts) or not all(vectors):
            self._semantic_error = "FastEmbed 返回的向量数量不完整"
            return None
        self._semantic_error = ""
        return vectors

    def build_semantic_index(
        self,
        *,
        document_id: str | None = None,
        force: bool = False,
    ) -> dict[str, Any]:
        """为本库当前分片建立向量索引；没有本地模型时只返回状态。"""
        query = """
            SELECT c.id, c.content, c.content_hash, c.stock_code, c.stock_name, c.model_name,
                   d.title, e.content_hash AS embedding_hash
            FROM chunks c
            JOIN documents d ON d.id = c.document_id
            LEFT JOIN embeddings e ON e.chunk_id = c.id AND e.model = ?
        """
        params: list[Any] = [self.semantic_model]
        if document_id:
            query += " WHERE c.document_id = ?"
            params.append(document_id)
        with self._lock:
            rows = self.connection.execute(query, params).fetchall()
        pending = [row for row in rows if force or row["embedding_hash"] != row["content_hash"]]
        cached = len(rows) - len(pending)
        total = len(rows)
        cache_hit_rate = cached / total if total else 0.0
        if not rows:
            return {
                "status": "empty",
                "indexed": 0,
                "cached": 0,
                "total": 0,
                "cache_hit_rate": 0.0,
                "model": self.semantic_model,
            }
        if not pending:
            return {
                "status": "ready",
                "indexed": 0,
                "cached": cached,
                "total": total,
                "cache_hit_rate": cache_hit_rate,
                "model": self.semantic_model,
            }
        vectors = self._embed(
            [
                "\n".join(
                    item
                    for item in (
                        row["title"], row["stock_code"], row["stock_name"], row["model_name"], row["content"],
                    )
                    if item
                )
                for row in pending
            ]
        )
        if vectors is None:
            return {
                "status": "skipped",
                "indexed": 0,
                "cached": cached,
                "total": total,
                "cache_hit_rate": cache_hit_rate,
                "model": self.semantic_model,
                "reason": self._semantic_error or "FastEmbed 本地模型不可用",
            }
        now = _now()
        with self._lock, self.connection:
            self.connection.executemany(
                """
                INSERT INTO embeddings(chunk_id, model, content_hash, dimensions, vector, updated_at)
                VALUES (?, ?, ?, ?, ?, ?)
                ON CONFLICT(chunk_id, model) DO UPDATE SET
                    content_hash=excluded.content_hash, dimensions=excluded.dimensions,
                    vector=excluded.vector, updated_at=excluded.updated_at
                """,
                [
                    (row["id"], self.semantic_model, row["content_hash"], len(vector), vector.tobytes(), now)
                    for row, vector in zip(pending, vectors)
                ],
            )
            self.connection.execute(
                "INSERT INTO knowledge_meta(key, value) VALUES('semantic_model', ?) "
                "ON CONFLICT(key) DO UPDATE SET value=excluded.value",
                (self.semantic_model,),
            )
            self.connection.execute(
                "INSERT INTO knowledge_meta(key, value) VALUES('semantic_indexed_at', ?) "
                "ON CONFLICT(key) DO UPDATE SET value=excluded.value",
                (now,),
            )
        return {
            "status": "ready",
            "indexed": len(vectors),
            "cached": cached,
            "total": total,
            "cache_hit_rate": cache_hit_rate,
            "model": self.semantic_model,
        }

    def semantic_status(self) -> dict[str, Any]:
        with self._lock:
            total_chunks = int(self.connection.execute("SELECT COUNT(*) FROM chunks").fetchone()[0])
            embedding_count = int(self.connection.execute(
                "SELECT COUNT(*) FROM embeddings WHERE model = ?", (self.semantic_model,)
            ).fetchone()[0])
            current_embeddings = int(self.connection.execute(
                """
                SELECT COUNT(*)
                FROM embeddings e JOIN chunks c ON c.id = e.chunk_id
                WHERE e.model = ? AND e.content_hash = c.content_hash
                """,
                (self.semantic_model,),
            ).fetchone()[0])
            stale_embeddings = int(self.connection.execute(
                """
                SELECT COUNT(*)
                FROM embeddings e JOIN chunks c ON c.id = e.chunk_id
                WHERE e.model = ? AND e.content_hash <> c.content_hash
                """,
                (self.semantic_model,),
            ).fetchone()[0])
        if total_chunks and current_embeddings == total_chunks and not stale_embeddings:
            status = "ready"
        elif embedding_count:
            status = "partial"
        else:
            status = "not_indexed"
        return {
            "status": status,
            "model": self.semantic_model,
            "embeddings": embedding_count,
            "total_chunks": total_chunks,
            "current_embeddings": current_embeddings,
            "stale_embeddings": stale_embeddings,
            "local_model_available": self._has_local_semantic_model(),
            "reason": self._semantic_error,
        }

    def get_document(self, document_id: str) -> dict[str, Any] | None:
        with self._lock:
            row = self.connection.execute(
                """
                SELECT d.*, COUNT(DISTINCT c.id) AS chunk_count,
                       COUNT(DISTINCT e.chunk_id) AS embedding_count
                FROM documents d
                LEFT JOIN chunks c ON c.document_id = d.id
                LEFT JOIN embeddings e ON e.chunk_id = c.id AND e.model = ?
                WHERE d.id = ?
                GROUP BY d.id
                """,
                (self.semantic_model, document_id),
            ).fetchone()
        if row is None:
            return None
        return self._result_document(row, chunks=row["chunk_count"], embeddings=row["embedding_count"])

    def list_documents(self, *, limit: int = 200) -> list[dict[str, Any]]:
        with self._lock:
            rows = self.connection.execute(
                """
                SELECT d.*, COUNT(DISTINCT c.id) AS chunk_count,
                       COUNT(DISTINCT e.chunk_id) AS embedding_count
                FROM documents d
                LEFT JOIN chunks c ON c.document_id = d.id
                LEFT JOIN embeddings e ON e.chunk_id = c.id AND e.model = ?
                GROUP BY d.id
                ORDER BY d.updated_at DESC, d.title COLLATE NOCASE
                LIMIT ?
                """,
                (self.semantic_model, max(1, min(int(limit), 1000))),
            ).fetchall()
        return [
            self._result_document(row, chunks=row["chunk_count"], embeddings=row["embedding_count"])
            for row in rows
        ]

    # documents 是路由层可读的短别名，保留 list_documents 的完整接口。
    documents = list_documents

    def list_chunks(self, document_id: str, *, limit: int = 200) -> list[dict[str, Any]]:
        with self._lock:
            rows = self.connection.execute(
                """
                SELECT id, document_id, chunk_index, content, metadata_json, tags_json,
                       stock_code, stock_name, model_name, parent_case_id, chunk_role,
                       mentioned_stocks_json, mentioned_attributes_json,
                       mentioned_highest_board, created_at
                FROM chunks WHERE document_id = ? ORDER BY chunk_index LIMIT ?
                """,
                (document_id, max(1, min(int(limit), 1000))),
            ).fetchall()
        return [
            {
                "chunk_id": str(row["id"]),
                "document_id": row["document_id"],
                "source_id": row["document_id"],
                "chunk_index": row["chunk_index"],
                "content": row["content"],
                "metadata": _json_loads(row["metadata_json"], {}),
                "tags": _normalize_tags(_json_loads(row["tags_json"], [])),
                "stock_code": row["stock_code"],
                "stock_name": row["stock_name"],
                "model_name": row["model_name"],
                "parent_case_id": row["parent_case_id"],
                "chunk_role": row["chunk_role"],
                "mentioned_stocks": _json_loads(row["mentioned_stocks_json"], []),
                "mentioned_attributes": _json_loads(row["mentioned_attributes_json"], []),
                "mentioned_highest_board": row["mentioned_highest_board"],
                "evidence_date": _case_date_range(
                    _json_loads(row["metadata_json"], {})
                ),
                "created_at": row["created_at"],
            }
            for row in rows
        ]

    def get_case(self, case_id: str) -> dict[str, Any] | None:
        with self._lock:
            row = self.connection.execute(
                "SELECT * FROM cases WHERE id = ?", (case_id,)
            ).fetchone()
        if row is None:
            return None
        data = dict(row)
        data["explicit_attributes"] = _json_loads(
            data.pop("explicit_attributes_json", "[]"), []
        )
        data["metadata"] = _json_loads(data.pop("metadata_json", "{}"), {})
        data["evidence_date"] = _case_date_range(data)
        return data

    def expand_case(self, case_id: str, *, chunk_limit: int = 3) -> dict[str, Any] | None:
        case = self.get_case(case_id)
        if case is None:
            return None
        with self._lock:
            rows = self.connection.execute(
                """
                SELECT id, chunk_index, content, chunk_role, metadata_json
                FROM chunks WHERE parent_case_id = ? ORDER BY chunk_index LIMIT ?
                """,
                (case_id, max(1, min(int(chunk_limit), 20))),
            ).fetchall()
        case["chunks"] = [
            {
                "chunk_id": str(row["id"]),
                "chunk_index": row["chunk_index"],
                "content": row["content"],
                "chunk_role": row["chunk_role"],
                "metadata": _json_loads(row["metadata_json"], {}),
            }
            for row in rows
        ]
        return case

    def list_tags(self) -> list[dict[str, Any]]:
        document_counts: defaultdict[str, set[str]] = defaultdict(set)
        chunk_counts: Counter[str] = Counter()
        with self._lock:
            rows = self.connection.execute("SELECT document_id, tags_json FROM chunks").fetchall()
        for row in rows:
            for tag in _normalize_tags(_json_loads(row["tags_json"], [])):
                document_counts[tag].add(row["document_id"])
                chunk_counts[tag] += 1
        return [
            {
                "tag": tag,
                "document_count": len(document_counts[tag]),
                "chunk_count": chunk_counts[tag],
            }
            for tag in sorted(document_counts, key=lambda item: (-len(document_counts[item]), item))
        ]

    tags = list_tags

    def update_document_tags(
        self,
        document_id: str,
        tags: Sequence[str] | str | None,
    ) -> dict[str, Any]:
        """替换一个来源的人工标签，同时保留该资料原有的案例标签。"""
        new_tags = _normalize_tags(tags)
        with self._lock, self.connection:
            document = self.connection.execute(
                "SELECT id, metadata_json FROM documents WHERE id = ?", (document_id,)
            ).fetchone()
            if document is None:
                raise KeyError(f"未找到历史模型资料：{document_id}")
            metadata = _json_loads(document["metadata_json"], {})
            all_document_tags = _normalize_tags([*new_tags, *_metadata_tags(metadata)])
            self.connection.execute(
                "UPDATE documents SET tags_json = ?, updated_at = ? WHERE id = ?",
                (_json_dumps(all_document_tags), _now(), document_id),
            )
            chunk_rows = self.connection.execute(
                "SELECT id, metadata_json FROM chunks WHERE document_id = ?", (document_id,)
            ).fetchall()
            self.connection.executemany(
                "UPDATE chunks SET tags_json = ?, search_text = ? WHERE id = ?",
                [
                    (
                        _json_dumps(_normalize_tags([*all_document_tags, *_metadata_tags(_json_loads(row["metadata_json"], {}))])),
                        self._fts_text(
                            title=self.connection.execute(
                                "SELECT title FROM documents WHERE id = ?", (document_id,)
                            ).fetchone()[0],
                            content=self.connection.execute(
                                "SELECT content FROM chunks WHERE id = ?", (row["id"],)
                            ).fetchone()[0],
                            metadata=_json_loads(row["metadata_json"], {}),
                            tags=_normalize_tags([*all_document_tags, *_metadata_tags(_json_loads(row["metadata_json"], {}))]),
                        ),
                        row["id"],
                    )
                    for row in chunk_rows
                ],
            )
            self._rebuild_fts_unlocked()
        result = self.get_document(document_id)
        assert result is not None
        return result

    update_source_tags = update_document_tags
    set_tags = update_document_tags

    def delete_document(self, document_id: str) -> bool:
        with self._lock, self.connection:
            cursor = self.connection.execute("DELETE FROM documents WHERE id = ?", (document_id,))
            self._rebuild_fts_unlocked()
        return bool(cursor.rowcount)

    def _fts_scores(self, terms: Sequence[str], *, limit: int) -> dict[int, float]:
        if not self.fts_available or not terms:
            return {}
        quoted = [f'"{term.replace(chr(34), chr(34) * 2)}"' for term in terms[:40] if term]
        if not quoted:
            return {}
        try:
            with self._lock:
                rows = self.connection.execute(
                    """
                    SELECT CAST(chunk_id AS INTEGER) AS chunk_id, bm25(chunks_fts) AS rank
                    FROM chunks_fts WHERE chunks_fts MATCH ? ORDER BY rank ASC LIMIT ?
                    """,
                    (" OR ".join(quoted), max(limit * 8, 30)),
                ).fetchall()
        except sqlite3.OperationalError:
            return {}
        raw = {int(row["chunk_id"]): -float(row["rank"]) for row in rows}
        if not raw:
            return {}
        minimum, maximum = min(raw.values()), max(raw.values())
        if maximum == minimum:
            return {key: 1.0 for key in raw}
        return {key: (value - minimum) / (maximum - minimum) for key, value in raw.items()}

    def _semantic_scores(self, query: str) -> dict[int, float]:
        with self._lock:
            rows = self.connection.execute(
                "SELECT chunk_id, dimensions, vector FROM embeddings WHERE model = ?",
                (self.semantic_model,),
            ).fetchall()
        if not rows:
            return {}
        query_vectors = self._embed([query])
        if not query_vectors:
            return {}
        query_vector = query_vectors[0]
        scores: dict[int, float] = {}
        for row in rows:
            vector = _vector_from_blob(row["vector"], row["dimensions"])
            if vector is not None:
                scores[int(row["chunk_id"])] = max(0.0, _dot(query_vector, vector))
        return scores

    @staticmethod
    def _matches_tags(row: Mapping[str, Any], tags: Sequence[str]) -> bool:
        if not tags:
            return True
        row_tags = set(_normalize_tags(_json_loads(str(row["tags_json"]), [])))
        return all(tag in row_tags for tag in tags)

    @staticmethod
    def _exact_score(
        row: Mapping[str, Any],
        *,
        query: str,
        stock_code: str | None,
        stock_name: str | None,
    ) -> float:
        query_codes = set(_STOCK_CODE_RE.findall(query))
        if stock_code:
            query_codes.add(str(stock_code).strip())
        row_code = str(row["stock_code"] or "").strip()
        score = 1.0 if row_code and row_code in query_codes else 0.0
        requested_name = _normalize_text(stock_name or "")
        row_name = _normalize_text(row["stock_name"] or "")
        normalized_query = _normalize_text(query)
        if row_name and (row_name == requested_name or row_name in normalized_query):
            score = max(score, 0.88)
        model_name = _normalize_text(row["model_name"] or "")
        if model_name and model_name in normalized_query:
            score = max(score, 0.65)
        return score

    def search(
        self,
        query: str,
        *,
        limit: int = 8,
        stock_code: str | None = None,
        stock_name: str | None = None,
        tags: Sequence[str] | str | None = None,
        semantic: bool = True,
    ) -> list[dict[str, Any]]:
        """股票代码/名称精确、FTS 和本地语义（可用时）混合检索。"""
        normalized_query = _normalize_text(query)
        requested_tags = _normalize_tags(tags)
        terms = _tokens(" ".join(part for part in (normalized_query, stock_code or "", stock_name or "") if part))
        if not (terms or requested_tags):
            return []
        maximum = max(1, min(int(limit), 30))
        with self._lock:
            rows = self.connection.execute(
                """
                SELECT c.*, d.title, d.source_path, d.source_type, d.updated_at
                FROM chunks c JOIN documents d ON d.id = c.document_id
                """
            ).fetchall()
        materialized = [row for row in rows if self._matches_tags(row, requested_tags)]
        if not materialized:
            return []
        fts_scores = self._fts_scores(terms, limit=maximum)
        query_counter = _counter(terms)
        lexical_scores = {
            int(row["id"]): _cosine(query_counter, _counter(str(row["search_text"]).split()))
            for row in materialized
        }
        exact_scores = {
            int(row["id"]): self._exact_score(
                row,
                query=normalized_query,
                stock_code=stock_code,
                stock_name=stock_name,
            )
            for row in materialized
        }
        semantic_scores = self._semantic_scores(normalized_query) if semantic and normalized_query else {}
        semantic_used = bool(semantic_scores)

        candidate_ids = {
            int(row["id"])
            for row in materialized
            if exact_scores[int(row["id"])] > 0 or fts_scores.get(int(row["id"]), 0) > 0
        }
        candidate_ids.update(
            chunk_id
            for chunk_id, _score in sorted(lexical_scores.items(), key=lambda item: item[1], reverse=True)[: maximum * 8]
            if _score > 0
        )
        candidate_ids.update(
            chunk_id
            for chunk_id, _score in sorted(semantic_scores.items(), key=lambda item: item[1], reverse=True)[: maximum * 8]
            if _score > 0
        )
        if not candidate_ids and requested_tags:
            candidate_ids = {int(row["id"]) for row in materialized}

        result_rows: list[dict[str, Any]] = []
        for row in materialized:
            chunk_id = int(row["id"])
            if chunk_id not in candidate_ids:
                continue
            exact_score = exact_scores.get(chunk_id, 0.0)
            fts_score = fts_scores.get(chunk_id, 0.0)
            lexical_score = lexical_scores.get(chunk_id, 0.0)
            semantic_score = semantic_scores.get(chunk_id, 0.0)
            if semantic_used:
                score = exact_score * 0.45 + fts_score * 0.22 + lexical_score * 0.13 + semantic_score * 0.20
            else:
                score = exact_score * 0.55 + fts_score * 0.30 + lexical_score * 0.15
            metadata = _json_loads(row["metadata_json"], {})
            item = {
                "chunk_id": str(chunk_id),
                "source_id": row["document_id"],
                "document_id": row["document_id"],
                "title": row["title"],
                "source_path": row["source_path"],
                "source_type": row["source_type"],
                "content": row["content"],
                "metadata": metadata if isinstance(metadata, dict) else {},
                "tags": _normalize_tags(_json_loads(row["tags_json"], [])),
                "stock_code": row["stock_code"],
                "stock_name": row["stock_name"],
                "model_name": row["model_name"],
                "parent_case_id": row["parent_case_id"],
                "chunk_role": row["chunk_role"],
                "mentioned_stocks": _json_loads(row["mentioned_stocks_json"], []),
                "mentioned_attributes": _json_loads(row["mentioned_attributes_json"], []),
                "mentioned_highest_board": row["mentioned_highest_board"],
                "evidence_date": _case_date_range(
                    metadata if isinstance(metadata, dict) else {}
                ),
                "score": round(score, 6),
                "retrieval_score": round(score, 6),
                "exact_score": round(exact_score, 6),
                "fts_score": round(fts_score, 6),
                "semantic_score": round(semantic_score, 6),
                "lexical_score": round(lexical_score, 6),
                "retrieval_mode": (
                    "股票代码/名称精确 + FTS + FastEmbed本地语义"
                    if semantic_used
                    else "股票代码/名称精确 + FTS + 轻量本地相似度"
                ),
                "updated_at": row["updated_at"],
            }
            result_rows.append(item)
        result_rows.sort(
            key=lambda item: (item["score"], item["exact_score"], item["updated_at"]), reverse=True
        )

        selected: list[dict[str, Any]] = []
        per_case: Counter[str] = Counter()
        for item in result_rows:
            case_key = item.get("parent_case_id") or f"chunk:{item['chunk_id']}"
            if per_case[case_key] >= 2:
                continue
            selected.append(item)
            per_case[case_key] += 1
            if len(selected) >= maximum:
                break
        return selected

    search_cases = search
    test_search = search

    def stats(self) -> dict[str, Any]:
        with self._lock:
            documents = self.connection.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
            chunks = self.connection.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
            cases = self.connection.execute("SELECT COUNT(*) FROM cases").fetchone()[0]
            embeddings = self.connection.execute(
                "SELECT COUNT(*) FROM embeddings WHERE model = ?", (self.semantic_model,)
            ).fetchone()[0]
        return {
            "database_path": str(self.path),
            "database_name": self.path.name,
            "documents": int(documents),
            "chunks": int(chunks),
            "cases": int(cases),
            "embeddings": int(embeddings),
            "fts_available": self.fts_available,
            "semantic": self.semantic_status(),
        }


# 便于后续模块按业务名称导入，不与现有 review_app.knowledge.KnowledgeStore 混淆。
HistoricalModelKnowledgeStore = DragonKnowledgeStore
DragonKnowledge = DragonKnowledgeStore


__all__ = [
    "DRAGON_KNOWLEDGE_DB_PATH",
    "SEMANTIC_MODEL",
    "ParsedUnit",
    "DragonKnowledgeStore",
    "DragonKnowledge",
    "HistoricalModelKnowledgeStore",
]

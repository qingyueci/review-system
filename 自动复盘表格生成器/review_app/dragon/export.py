"""将已持久化的首板分析保存为可阅读的 Word 文件。"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document

from .schemas import DragonAnalysisRecord


def _append_model_value(document: Document, value: Any, *, level: int = 0) -> None:
    if isinstance(value, dict):
        for key, nested in value.items():
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = None
            paragraph.add_run(str(key)).bold = True
            if isinstance(nested, (dict, list)):
                _append_model_value(document, nested, level=level + 1)
            else:
                paragraph.add_run(f"：{nested}")
        return
    if isinstance(value, list):
        for item in value:
            if isinstance(item, (dict, list)):
                _append_model_value(document, item, level=level + 1)
            else:
                document.add_paragraph(str(item), style="List Bullet")
        return
    document.add_paragraph(str(value))


def build_dragon_analysis_docx(records: list[DragonAnalysisRecord]) -> bytes:
    if not records:
        raise ValueError("没有可保存的首板分析结果")

    first = records[0]
    document = Document()
    document.add_heading(f"首板布局分析 {first.trade_date.isoformat()}", level=0)
    document.add_paragraph(f"候选数量：{len(records)}")

    snapshot = first.context.review_snapshot
    if snapshot:
        document.add_heading("当日确认复盘", level=1)
        snapshot_items = [
            ("周期阶段", snapshot.period_stage),
            ("市场核心", snapshot.market_core),
            ("超预期点", "；".join(snapshot.positive_surprises)),
            ("负反馈", "；".join(snapshot.negative_feedback)),
            ("有效方向", "；".join(snapshot.effective_directions)),
            ("明日布局任务", "；".join(snapshot.layout_tasks)),
            ("失效条件", "；".join(snapshot.failure_conditions)),
            ("用户补充", snapshot.user_notes),
        ]
        for label, text in snapshot_items:
            if text:
                paragraph = document.add_paragraph()
                paragraph.add_run(f"{label}：").bold = True
                paragraph.add_run(text)

    document.add_heading("候选结论", level=1)
    for record in records:
        result = record.result
        document.add_heading(
            f"{record.stock_name}（{record.stock_code}）· {result.conclusion}",
            level=2,
        )
        document.add_paragraph("基础标准：" + ("通过" if record.basic_pass else "不通过"))
        exclusion_reason = str(result.model_output.get("exclusion_reason") or "").strip()
        if exclusion_reason:
            document.add_paragraph(exclusion_reason)
        if result.analysis:
            _append_model_value(document, result.analysis)
        if result.history_dates:
            document.add_paragraph("、".join(result.history_dates))

        checks = record.context.screening.checks
        if checks:
            table = document.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for cell, label in zip(table.rows[0].cells, ("检查项", "实际", "标准", "结果")):
                cell.text = label
            for check in checks:
                row = table.add_row().cells
                row[0].text = check.rule_name
                row[1].text = str(check.actual_value)
                row[2].text = str(check.threshold)
                row[3].text = check.status

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


__all__ = ["build_dragon_analysis_docx"]

from datetime import date
from io import BytesIO
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _set_chinese_font(run, name: str = "微软雅黑") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    run.add_text(" 页")


def _style_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    normal = document.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "微软雅黑")
    for style_name, size, color in (
        ("Title", 22, "1F4E78"),
        ("Heading 1", 15, "1F4E78"),
        ("Heading 2", 12, "365F91"),
    ):
        style = document.styles[style_name]
        style.font.name = "微软雅黑"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "微软雅黑")


def _add_markdown(document: Document, analysis: str) -> None:
    for raw_line in analysis.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 2)
            document.add_heading(heading.group(2).strip(), level=level)
            continue
        if re.match(r"^[-*]\s+", line):
            paragraph = document.add_paragraph(style="List Bullet")
            text = re.sub(r"^[-*]\s+", "", line)
        elif re.match(r"^\d+[.、]\s*", line):
            paragraph = document.add_paragraph(style="List Number")
            text = re.sub(r"^\d+[.、]\s*", "", line)
        else:
            paragraph = document.add_paragraph()
            text = line
        run = paragraph.add_run(re.sub(r"\*\*(.*?)\*\*", r"\1", text))
        _set_chinese_font(run)


def generate_analysis_docx(
    analysis: str,
    sources: list[dict],
    *,
    review_date: str | None = None,
) -> tuple[bytes, str]:
    document = Document()
    _style_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("刺大公开框架辅助复盘")
    _set_chinese_font(run)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"复盘日期：{review_date or date.today().isoformat()}")
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    _set_chinese_font(subtitle_run)

    notice = document.add_paragraph()
    notice_run = notice.add_run(
        "说明：本文由公开历史资料检索辅助生成，不代表原作者本人观点，不构成投资建议。"
    )
    notice_run.italic = True
    notice_run.font.color.rgb = RGBColor(127, 127, 127)
    _set_chinese_font(notice_run)

    _add_markdown(document, analysis)
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("引用的公开历史资料", level=1)
    labels = {
        "qa": "刺大公开回复",
        "community": "社区精选评论，仅作辅助",
        "manual": "人工整理体系",
        "post": "复盘主帖",
    }
    for index, source in enumerate(sources, 1):
        paragraph = document.add_paragraph(style="List Number")
        source_type = labels.get(source["source_type"], "公开资料")
        label = f"{source['title']}（{source['published_at'][:10]}，{source_type}）"
        run = paragraph.add_run(label + " ")
        _set_chinese_font(run)
        _add_hyperlink(paragraph, f"[资料{index} 原文]", source["source_url"])

    for section in document.sections:
        _add_page_number(section.footer.paragraphs[0])

    output = BytesIO()
    document.save(output)
    safe_date = re.sub(r'[<>:"/\\|?*]', "_", review_date or date.today().isoformat())
    return output.getvalue(), f"刺大框架复盘分析_{safe_date}.docx"

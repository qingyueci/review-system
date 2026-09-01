from collections import Counter
from datetime import datetime
import json
import math
import re
import sqlite3
from typing import Callable

from docx import Document
import jieba
import numpy as np

from .cleaning import content_hash, normalize_text, split_chunks
from .config import CRAWL_TOP_POST_LIMIT, KNOWLEDGE_DB_PATH, MANUAL_SYSTEM_DOCX
from .crawler import TgbCrawler

jieba.setLogLevel(20)

STOP_WORDS = {
    "一个", "今天", "明天", "这个", "那个", "还是", "就是", "感觉", "怎么", "什么",
    "已经", "没有", "可以", "如果", "因为", "所以", "但是", "然后", "一下", "现在",
    "市场", "个股", "板块", "老师", "刺大", "发财",
}
CRAWL_PIPELINE_VERSION = 3
SEMANTIC_MODEL = "BAAI/bge-small-zh-v1.5"
SEMANTIC_INDEX_VERSION = "answer-v1"
SEMANTIC_SAME_POST_THRESHOLD = 0.91
SEMANTIC_CROSS_POST_THRESHOLD = 0.94
SOURCE_WEIGHTS = {
    "qa": 1.0,
    "post": 0.9,
    "manual": 0.82,
    "community": 0.55,
}
SCOPE_WEIGHTS = {
    "top_year": 1.0,
    "recent_qa": 1.0,
    "recent_archive": 0.72,
    "manual": 1.0,
}
CONCEPT_TERMS = {
    "首板出身": ("首板", "出身", "发酵来源", "起点"),
    "个股任务": ("任务", "角色", "使命", "负责"),
    "布局关系": ("布局", "协同", "压制", "反推", "带动"),
    "市场地位": ("地位", "辨识度", "主动性", "独立性", "身位"),
    "竞价确认": ("竞价", "确认", "预期", "开盘"),
    "失败条件": ("失败", "失效", "证伪", "不及预期"),
}

_EMBEDDING_MODELS: dict[str, object] = {}
_OPPOSITE_TERMS = (
    ("上涨", "下跌"), ("涨", "跌"), ("买", "卖"), ("强", "弱"),
    ("看多", "看空"), ("成功", "失败"), ("加仓", "减仓"),
)
_NEGATION_TERMS = ("不是", "没有", "不能", "不应", "别", "未", "无", "避免", "禁止")


def _tokens(value: str, *, maximum: int = 80) -> list[str]:
    text = normalize_text(value)
    words = [
        word.strip().lower()
        for word in jieba.lcut(text)
        if word.strip() and (len(word.strip()) > 1 or word.strip().isascii())
    ]
    words = [word for word in words if word not in STOP_WORDS and not re.fullmatch(r"\W+", word)]
    cjk = re.sub(r"[^\u4e00-\u9fff]", "", text)
    bigrams = [f"g_{cjk[index:index + 2]}" for index in range(max(0, len(cjk) - 1))]
    ranked = [item for item, _count in Counter(words + bigrams).most_common(maximum)]
    return ranked


def _index_text(value: str) -> str:
    return " ".join(_tokens(value, maximum=400))


def _vector_features(value: str) -> Counter:
    """构造轻量本地语义向量，不调用外部模型、不额外消耗额度。"""
    normalized = normalize_text(value).lower()
    features: Counter = Counter(_tokens(normalized, maximum=160))
    for concept, aliases in CONCEPT_TERMS.items():
        if any(alias in normalized for alias in aliases):
            features[f"concept:{concept}"] += 3
    return features


def _stored_vector_features(search_text: str, value: str) -> Counter:
    """复用入库时的分词结果，避免每次检索重新切分全部知识库。"""
    features: Counter = Counter(search_text.split())
    normalized = normalize_text(value).lower()
    for concept, aliases in CONCEPT_TERMS.items():
        if any(alias in normalized for alias in aliases):
            features[f"concept:{concept}"] += 3
    return features


def _cosine(left: Counter, right: Counter) -> float:
    if not left or not right:
        return 0.0
    common = set(left) & set(right)
    numerator = sum(left[item] * right[item] for item in common)
    left_norm = math.sqrt(sum(value * value for value in left.values()))
    right_norm = math.sqrt(sum(value * value for value in right.values()))
    return numerator / (left_norm * right_norm) if left_norm and right_norm else 0.0


def _time_score(value: str, source_type: str) -> float:
    if source_type == "manual":
        return 0.85
    try:
        published = datetime.fromisoformat(value.replace("Z", "+00:00")).replace(
            tzinfo=None
        )
    except (TypeError, ValueError):
        return 0.45
    age_days = max(0, (datetime.now() - published).days)
    return max(0.2, math.exp(-age_days / 730))


def _topic_score(query_features: Counter, content_features: Counter) -> float:
    query_concepts = {key for key in query_features if key.startswith("concept:")}
    if not query_concepts:
        return 0.0
    return len(query_concepts & set(content_features)) / len(query_concepts)


def _embedding_model(model_name: str = SEMANTIC_MODEL):
    """按需载入本地 ONNX 模型；平时检索不启动模型。"""
    if model_name not in _EMBEDDING_MODELS:
        from fastembed import TextEmbedding

        cache_dir = KNOWLEDGE_DB_PATH.parent / "fastembed_cache"
        cache_dir.mkdir(parents=True, exist_ok=True)
        _EMBEDDING_MODELS[model_name] = TextEmbedding(
            model_name=model_name,
            cache_dir=str(cache_dir),
        )
    return _EMBEDDING_MODELS[model_name]


def _semantic_text(row: sqlite3.Row | dict) -> str:
    return normalize_text(row["answer"])


def _semantically_compatible(left: str, right: str) -> bool:
    """数字、否定和相反方向不一致时禁止合并。"""
    left_numbers = set(re.findall(r"\d+(?:\.\d+)?%?", left))
    right_numbers = set(re.findall(r"\d+(?:\.\d+)?%?", right))
    if left_numbers != right_numbers and (left_numbers or right_numbers):
        return False
    for term in _NEGATION_TERMS:
        if (term in left) != (term in right):
            return False
    for positive, negative in _OPPOSITE_TERMS:
        if (positive in left and negative in right) or (negative in left and positive in right):
            return False
    return True


class KnowledgeStore:
    def __init__(self, path=KNOWLEDGE_DB_PATH) -> None:
        self.path = path
        self.path.parent.mkdir(parents=True, exist_ok=True)
        self.connection = sqlite3.connect(self.path)
        self.connection.row_factory = sqlite3.Row
        self._create_schema()

    def close(self) -> None:
        self.connection.close()

    def __enter__(self):
        return self

    def __exit__(self, *_args) -> None:
        self.close()

    def _create_schema(self) -> None:
        self.connection.executescript("""
            PRAGMA journal_mode=WAL;
            CREATE TABLE IF NOT EXISTS posts (
                url TEXT PRIMARY KEY,
                topic_id TEXT,
                title TEXT NOT NULL,
                published_at TEXT NOT NULL,
                views INTEGER NOT NULL DEFAULT 0,
                reply_count INTEGER NOT NULL DEFAULT 0,
                likes INTEGER NOT NULL DEFAULT 0,
                summary TEXT NOT NULL DEFAULT '',
                body TEXT NOT NULL DEFAULT '',
                body_hash TEXT NOT NULL DEFAULT '',
                body_truncated INTEGER NOT NULL DEFAULT 0,
                comments_accessible INTEGER NOT NULL DEFAULT 0,
                useful_comment_count INTEGER NOT NULL DEFAULT 0,
                total_comment_pages INTEGER NOT NULL DEFAULT 0,
                scanned_comment_pages TEXT NOT NULL DEFAULT '[]',
                scope TEXT NOT NULL DEFAULT 'top_year',
                capture_mode TEXT NOT NULL DEFAULT 'public_http',
                pipeline_version INTEGER NOT NULL DEFAULT 1,
                updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS qa_pairs (
                reply_id TEXT PRIMARY KEY,
                post_url TEXT NOT NULL,
                question TEXT NOT NULL DEFAULT '',
                answer TEXT NOT NULL,
                question_author TEXT NOT NULL DEFAULT '',
                published_at TEXT NOT NULL DEFAULT '',
                floor INTEGER NOT NULL DEFAULT 0,
                likes INTEGER NOT NULL DEFAULT 0,
                source_url TEXT NOT NULL,
                content_hash TEXT NOT NULL,
                is_retrievable INTEGER NOT NULL DEFAULT 1,
                duplicate_of TEXT NOT NULL DEFAULT '',
                similarity_score REAL NOT NULL DEFAULT 0,
                semantic_group_id TEXT NOT NULL DEFAULT '',
                dedupe_reason TEXT NOT NULL DEFAULT '',
                FOREIGN KEY(post_url) REFERENCES posts(url) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS chunks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                source_type TEXT NOT NULL,
                source_id TEXT NOT NULL,
                post_url TEXT NOT NULL,
                title TEXT NOT NULL,
                published_at TEXT NOT NULL,
                content TEXT NOT NULL,
                source_url TEXT NOT NULL,
                chunk_order INTEGER NOT NULL DEFAULT 0,
                search_text TEXT NOT NULL,
                UNIQUE(source_type, source_id, chunk_order)
            );
            CREATE TABLE IF NOT EXISTS crawl_runs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                started_at TEXT NOT NULL,
                finished_at TEXT NOT NULL,
                discovered INTEGER NOT NULL,
                fetched INTEGER NOT NULL,
                reused INTEGER NOT NULL,
                failed INTEGER NOT NULL,
                error_details TEXT NOT NULL DEFAULT ''
            );
            CREATE TABLE IF NOT EXISTS local_sources (
                path TEXT PRIMARY KEY,
                title TEXT NOT NULL,
                modified_at TEXT NOT NULL,
                content_hash TEXT NOT NULL,
                content TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS qa_embeddings (
                reply_id TEXT PRIMARY KEY,
                model TEXT NOT NULL,
                content_hash TEXT NOT NULL,
                dimensions INTEGER NOT NULL,
                vector BLOB NOT NULL,
                updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS knowledge_meta (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL
            );
        """)
        existing_columns = {
            row["name"] for row in self.connection.execute("PRAGMA table_info(posts)").fetchall()
        }
        for name in ("body_truncated", "comments_accessible", "useful_comment_count"):
            if name not in existing_columns:
                self.connection.execute(
                    f"ALTER TABLE posts ADD COLUMN {name} INTEGER NOT NULL DEFAULT 0"
                )
        if "pipeline_version" not in existing_columns:
            self.connection.execute(
                "ALTER TABLE posts ADD COLUMN pipeline_version INTEGER NOT NULL DEFAULT 1"
            )
        if "capture_mode" not in existing_columns:
            self.connection.execute(
                "ALTER TABLE posts ADD COLUMN capture_mode TEXT NOT NULL DEFAULT 'public_http'"
            )
        qa_columns = {
            row["name"] for row in self.connection.execute("PRAGMA table_info(qa_pairs)").fetchall()
        }
        qa_migrations = {
            "is_retrievable": "INTEGER NOT NULL DEFAULT 1",
            "duplicate_of": "TEXT NOT NULL DEFAULT ''",
            "similarity_score": "REAL NOT NULL DEFAULT 0",
            "semantic_group_id": "TEXT NOT NULL DEFAULT ''",
            "dedupe_reason": "TEXT NOT NULL DEFAULT ''",
        }
        for name, definition in qa_migrations.items():
            if name not in qa_columns:
                self.connection.execute(
                    f"ALTER TABLE qa_pairs ADD COLUMN {name} {definition}"
                )
        try:
            self.connection.execute("""
                CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts
                USING fts5(chunk_id UNINDEXED, search_text)
            """)
        except sqlite3.OperationalError as exc:
            raise RuntimeError("当前 Python 的 SQLite 不支持 FTS5，无法建立本地检索索引") from exc
        self.connection.commit()

    def post_state(self, url: str) -> dict | None:
        row = self.connection.execute(
            """
            SELECT reply_count, body_hash, pipeline_version, capture_mode
            FROM posts WHERE url = ?
            """,
            (url,),
        ).fetchone()
        return dict(row) if row else None

    def scope_urls(self, scope: str) -> list[str]:
        """读取指定分组，核心库存在后不再重新扫描和改写。"""
        return [
            row["url"]
            for row in self.connection.execute(
                "SELECT url FROM posts WHERE scope = ?", (scope,)
            ).fetchall()
        ]

    def update_post_metrics(self, post: dict, scope: str) -> None:
        self.connection.execute(
            """
            UPDATE posts
            SET views = ?, reply_count = ?, likes = ?, summary = ?, scope = ?, updated_at = ?
            WHERE url = ?
            """,
            (
                post["views"], post["reply_count"], post["likes"], post.get("summary", ""),
                scope, datetime.now().isoformat(timespec="seconds"), post["url"],
            ),
        )
        self.connection.commit()

    def upsert_post(
        self,
        post: dict,
        *,
        scope: str = "top_year",
        rebuild_index: bool = True,
    ) -> None:
        now = datetime.now().isoformat(timespec="seconds")
        with self.connection:
            self.connection.execute(
                """
                INSERT INTO posts (
                    url, topic_id, title, published_at, views, reply_count, likes, summary,
                    body, body_hash, total_comment_pages, scanned_comment_pages, scope, updated_at,
                    body_truncated, comments_accessible, useful_comment_count,
                    capture_mode, pipeline_version
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(url) DO UPDATE SET
                    topic_id=excluded.topic_id, title=excluded.title,
                    published_at=excluded.published_at, views=excluded.views,
                    reply_count=excluded.reply_count, likes=excluded.likes,
                    summary=excluded.summary, body=excluded.body,
                    body_hash=excluded.body_hash,
                    body_truncated=excluded.body_truncated,
                    comments_accessible=excluded.comments_accessible,
                    useful_comment_count=excluded.useful_comment_count,
                    capture_mode=excluded.capture_mode,
                    total_comment_pages=excluded.total_comment_pages,
                    scanned_comment_pages=excluded.scanned_comment_pages,
                    scope=excluded.scope, pipeline_version=excluded.pipeline_version,
                    updated_at=excluded.updated_at
                """,
                (
                    post["url"], post.get("topic_id", ""), post["title"], post["published_at"],
                    post["views"], post["reply_count"], post["likes"], post.get("summary", ""),
                    post["body"], post["body_hash"], post.get("total_comment_pages", 0),
                    json.dumps(post.get("scanned_comment_pages", []), ensure_ascii=False),
                    scope, now, int(post.get("body_truncated", False)),
                    int(post.get("comments_accessible", False)),
                    int(post.get("useful_comment_count", 0)),
                    post.get("capture_mode", "public_http"),
                    CRAWL_PIPELINE_VERSION,
                ),
            )
            self.connection.execute(
                "DELETE FROM qa_embeddings WHERE reply_id IN "
                "(SELECT reply_id FROM qa_pairs WHERE post_url = ?)",
                (post["url"],),
            )
            self.connection.execute("DELETE FROM qa_pairs WHERE post_url = ?", (post["url"],))
            self.connection.execute("DELETE FROM chunks WHERE post_url = ?", (post["url"],))

            for reply in post.get("author_replies", []):
                answer = normalize_text(reply["answer"])
                question = normalize_text(reply.get("question", ""))
                self.connection.execute(
                    """
                    INSERT OR REPLACE INTO qa_pairs (
                        reply_id, post_url, question, answer, question_author,
                        published_at, floor, likes, source_url, content_hash
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        reply["reply_id"], post["url"], question, answer,
                        reply.get("question_author", ""), reply.get("published_at", ""),
                        reply.get("floor", 0), reply.get("likes", 0), reply["source_url"],
                        content_hash(question, answer),
                    ),
                )

            article_prefix = "【网页公开节选】\n" if post.get("body_truncated") else ""
            for order, chunk in enumerate(split_chunks(post["body"])):
                self._insert_chunk(
                    "post", post["url"], post, article_prefix + chunk, post["url"], order,
                )
            for reply in post.get("author_replies", []):
                question = normalize_text(reply.get("question", ""))
                answer = normalize_text(reply["answer"])
                combined = f"用户问题：{question}\n刺大回复：{answer}" if question else f"刺大评论：{answer}"
                self._insert_chunk(
                    "qa", reply["reply_id"], post, combined, reply["source_url"], 0,
                )
            for comment in post.get("community_comments", []):
                answer = normalize_text(comment["answer"])
                quote = normalize_text(comment.get("question", ""))
                author = normalize_text(comment.get("author_name", "")) or "社区用户"
                likes = int(comment.get("likes", 0))
                combined = f"社区精选评论（{author}，点赞{likes}）：{answer}"
                if quote:
                    combined += f"\n引用上下文：{quote}"
                self._insert_chunk(
                    "community", comment["reply_id"], post, combined,
                    comment["source_url"], 0,
                )
        if rebuild_index:
            self.rebuild_fts()

    def import_manual_docx(self, path=MANUAL_SYSTEM_DOCX) -> dict:
        """把人工整理的 Word 体系文件清洗、切片后加入检索库。"""
        source_path = path.resolve()
        if not source_path.is_file():
            raise RuntimeError(f"没有找到本地体系文件：{source_path}")

        document = Document(source_path)
        sections: list[str] = []
        for paragraph in document.paragraphs:
            text = normalize_text(paragraph.text)
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.startswith("Heading") or style_name.startswith("标题"):
                text = f"章节：{text}"
            sections.append(text)
        for table in document.tables:
            for row in table.rows:
                values = [normalize_text(cell.text) for cell in row.cells]
                values = [value for value in values if value]
                if values:
                    sections.append(" | ".join(values))
        content = normalize_text("\n\n".join(sections))
        if not content:
            raise RuntimeError(f"本地体系文件没有可读取的文字：{source_path}")

        digest = content_hash(content)
        path_text = str(source_path)
        current = self.connection.execute(
            "SELECT content_hash FROM local_sources WHERE path = ?", (path_text,)
        ).fetchone()
        if current and current["content_hash"] == digest:
            return {"imported": False, "path": path_text, "chunks": 0}

        modified_at = datetime.fromtimestamp(source_path.stat().st_mtime).isoformat(
            timespec="seconds"
        )
        source_key = f"manual:{path_text}"
        source = {
            "url": source_key,
            "title": source_path.stem,
            "published_at": modified_at,
        }
        chunks = split_chunks(content)
        with self.connection:
            self.connection.execute(
                "DELETE FROM chunks WHERE source_type='manual' AND post_url = ?",
                (source_key,),
            )
            self.connection.execute(
                """
                INSERT INTO local_sources (
                    path, title, modified_at, content_hash, content, updated_at
                ) VALUES (?, ?, ?, ?, ?, ?)
                ON CONFLICT(path) DO UPDATE SET
                    title=excluded.title, modified_at=excluded.modified_at,
                    content_hash=excluded.content_hash, content=excluded.content,
                    updated_at=excluded.updated_at
                """,
                (
                    path_text, source_path.stem, modified_at, digest, content,
                    datetime.now().isoformat(timespec="seconds"),
                ),
            )
            for order, chunk in enumerate(chunks):
                self._insert_chunk(
                    "manual", source_key, source, chunk, source_path.as_uri(), order,
                )
        self.rebuild_fts()
        return {"imported": True, "path": path_text, "chunks": len(chunks)}

    def _insert_chunk(
        self,
        source_type: str,
        source_id: str,
        post: dict,
        content: str,
        source_url: str,
        order: int,
    ) -> None:
        self.connection.execute(
            """
            INSERT OR REPLACE INTO chunks (
                source_type, source_id, post_url, title, published_at,
                content, source_url, chunk_order, search_text
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                source_type, source_id, post["url"], post["title"], post["published_at"],
                content, source_url, order, _index_text(content),
            ),
        )

    def archive_recent_except(self, active_urls: list[str]) -> int:
        """把移出近期窗口的帖子降级归档，保留正文、回复和检索片段。"""
        if not active_urls:
            return 0
        placeholders = ",".join("?" for _ in active_urls)
        with self.connection:
            cursor = self.connection.execute(
                f"""
                UPDATE posts
                SET scope = 'recent_archive', updated_at = ?
                WHERE scope = 'recent_qa' AND url NOT IN ({placeholders})
                """,
                [datetime.now().isoformat(timespec="seconds"), *active_urls],
            )
        return cursor.rowcount

    def rebuild_fts(self) -> None:
        with self.connection:
            self.connection.execute("DELETE FROM chunks_fts")
            self.connection.execute(
                "INSERT INTO chunks_fts(chunk_id, search_text) SELECT id, search_text FROM chunks"
            )

    def search(self, query: str, *, limit: int = 12) -> list[dict]:
        terms = _tokens(query, maximum=30)
        if not terms:
            return []
        match_query = " OR ".join(f'"{term.replace(chr(34), chr(34) * 2)}"' for term in terms)
        keyword_rows = self.connection.execute(
            """
            SELECT c.*, bm25(chunks_fts) AS rank
            FROM chunks_fts
            JOIN chunks c ON c.id = CAST(chunks_fts.chunk_id AS INTEGER)
            WHERE chunks_fts MATCH ?
            ORDER BY rank ASC, c.published_at DESC
            LIMIT ?
            """,
            (match_query, limit * 6),
        ).fetchall()
        all_rows = [
            dict(row)
            for row in self.connection.execute(
                """
                SELECT c.*, COALESCE(p.scope, 'manual') AS scope
                FROM chunks c
                LEFT JOIN posts p ON p.url = c.post_url
                """
            ).fetchall()
        ]
        query_vector = _vector_features(query)
        vector_scores = {
            row["id"]: _cosine(
                query_vector,
                _stored_vector_features(
                    row["search_text"],
                    f"{row['title']} {row['content']}",
                ),
            )
            for row in all_rows
        }
        vector_ids = {
            item[0]
            for item in sorted(
                vector_scores.items(),
                key=lambda item: item[1],
                reverse=True,
            )[: limit * 6]
            if item[1] > 0
        }
        keyword_values = {
            row["id"]: -float(row["rank"])
            for row in keyword_rows
        }
        keyword_min = min(keyword_values.values(), default=0.0)
        keyword_max = max(keyword_values.values(), default=0.0)

        candidates = {
            row["id"]: row
            for row in all_rows
            if row["id"] in vector_ids or row["id"] in keyword_values
        }
        scored_rows = []
        for row in candidates.values():
            raw_keyword = keyword_values.get(row["id"], keyword_min)
            keyword_score = (
                (raw_keyword - keyword_min) / (keyword_max - keyword_min)
                if keyword_max > keyword_min
                else (1.0 if row["id"] in keyword_values else 0.0)
            )
            content_vector = _stored_vector_features(
                row["search_text"],
                f"{row['title']} {row['content']}",
            )
            vector_score = vector_scores.get(row["id"], 0.0)
            source_weight = (
                SOURCE_WEIGHTS.get(row["source_type"], 0.5)
                * SCOPE_WEIGHTS.get(row["scope"], 1.0)
            )
            time_score = _time_score(row["published_at"], row["source_type"])
            topic_score = _topic_score(query_vector, content_vector)
            final_score = (
                keyword_score * 0.30
                + vector_score * 0.30
                + source_weight * 0.22
                + topic_score * 0.11
                + time_score * 0.07
            )
            row.update(
                keyword_score=round(keyword_score, 4),
                vector_score=round(vector_score, 4),
                source_weight=source_weight,
                topic_score=round(topic_score, 4),
                time_score=round(time_score, 4),
                retrieval_score=round(final_score, 4),
                retrieval_mode="关键词 + 本地向量 + 来源/时间/题材权重",
            )
            scored_rows.append(row)
        scored_rows.sort(
            key=lambda row: (
                row["retrieval_score"],
                row["published_at"],
            ),
            reverse=True,
        )

        # 保证作者原文和人工体系有足够席位，再用社区观点补充。
        source_limits = {"qa": 4, "post": 3, "manual": 3, "community": 2}
        # 同一来源最多保留三条，避免单篇长文垄断上下文。
        selected: list[dict] = []
        per_post: Counter = Counter()
        per_type: Counter = Counter()
        for row in scored_rows:
            if per_post[row["post_url"]] >= 3:
                continue
            if per_type[row["source_type"]] >= source_limits.get(row["source_type"], 2):
                continue
            selected.append(row)
            per_post[row["post_url"]] += 1
            per_type[row["source_type"]] += 1
            if len(selected) >= limit:
                break
        # 配额没有填满时，按混合得分补齐。
        selected_ids = {item["id"] for item in selected}
        for row in scored_rows:
            if len(selected) >= limit:
                break
            if row["id"] in selected_ids or per_post[row["post_url"]] >= 3:
                continue
            selected.append(row)
            selected_ids.add(row["id"])
            per_post[row["post_url"]] += 1
        return selected

    def semantic_clean_qa(
        self,
        progress: Callable[[str, int, int], None] | None = None,
        *,
        model_name: str = SEMANTIC_MODEL,
    ) -> dict:
        """保留全部原始回复，只让每个语义组的代表回复进入检索层。"""
        rows = self.connection.execute(
            """
            SELECT q.*, p.title AS post_title, p.published_at AS post_published_at
            FROM qa_pairs q
            JOIN posts p ON p.url = q.post_url
            ORDER BY q.published_at DESC, q.floor DESC
            """
        ).fetchall()
        total = len(rows)
        cleaned_at = datetime.now().isoformat(timespec="seconds")
        cache_model = f"{model_name}:{SEMANTIC_INDEX_VERSION}"
        if not rows:
            with self.connection:
                self.connection.execute(
                    "INSERT OR REPLACE INTO knowledge_meta(key, value) VALUES('semantic_model', ?)",
                    (model_name,),
                )
                self.connection.execute(
                    "INSERT OR REPLACE INTO knowledge_meta(key, value) VALUES('semantic_cleaned_at', ?)",
                    (cleaned_at,),
                )
            return {
                "original": 0,
                "retrievable": 0,
                "duplicates": 0,
                "model": model_name,
                "cleaned_at": cleaned_at,
            }

        cached = {
            row["reply_id"]: row
            for row in self.connection.execute(
                "SELECT reply_id, content_hash, dimensions, vector FROM qa_embeddings WHERE model = ?",
                (cache_model,),
            ).fetchall()
        }
        vectors: list[np.ndarray | None] = [None] * total
        missing_indices: list[int] = []
        for index, row in enumerate(rows):
            stored = cached.get(row["reply_id"])
            if stored and stored["content_hash"] == row["content_hash"]:
                vector = np.frombuffer(stored["vector"], dtype=np.float32)
                if vector.size == stored["dimensions"]:
                    vectors[index] = vector
                    continue
            missing_indices.append(index)

        if missing_indices:
            if progress:
                progress("正在载入本地语义模型", 0, total)
            model = _embedding_model(model_name)
            texts = [_semantic_text(rows[index]) for index in missing_indices]
            generated = model.embed(texts, batch_size=64)
            inserts = []
            for completed, (index, vector) in enumerate(
                zip(missing_indices, generated, strict=True), 1
            ):
                normalized = np.asarray(vector, dtype=np.float32)
                norm = float(np.linalg.norm(normalized))
                if norm:
                    normalized /= norm
                vectors[index] = normalized
                row = rows[index]
                inserts.append(
                    (
                        row["reply_id"], cache_model, row["content_hash"],
                        int(normalized.size), normalized.tobytes(), cleaned_at,
                    )
                )
                if progress and (completed % 256 == 0 or completed == len(missing_indices)):
                    progress(
                        f"正在生成本地语义索引 {completed}/{len(missing_indices)}",
                        completed,
                        len(missing_indices),
                    )
            with self.connection:
                self.connection.executemany(
                    """
                    INSERT OR REPLACE INTO qa_embeddings(
                        reply_id, model, content_hash, dimensions, vector, updated_at
                    ) VALUES (?, ?, ?, ?, ?, ?)
                    """,
                    inserts,
                )

        matrix = np.vstack([vector for vector in vectors if vector is not None]).astype(
            np.float32, copy=False
        )
        if len(matrix) != total:
            raise RuntimeError("本地语义向量生成不完整，请重新执行")

        def quality(index: int) -> tuple[float, str, int]:
            row = rows[index]
            score = len(normalize_text(row["answer"])) + min(int(row["likes"]), 30) * 8
            score += min(len(normalize_text(row["question"])), 120) * 0.25
            return score, row["published_at"], int(row["floor"])

        order = sorted(range(total), key=quality, reverse=True)
        representatives = np.empty_like(matrix)
        representative_indices: list[int] = []
        representative_count = 0
        assignments: list[tuple[int, str, float, str, str] | None] = [None] * total
        exact_representatives: dict[str, int] = {}

        for completed, index in enumerate(order, 1):
            row = rows[index]
            text = _semantic_text(row)
            matched_rep_position: int | None = None
            similarity = 0.0
            reason = "代表回复"
            exact_position = exact_representatives.get(row["content_hash"])
            if exact_position is not None:
                matched_rep_position = exact_position
                similarity = 1.0
                reason = "文本重复"
            elif representative_count:
                scores = representatives[:representative_count] @ matrix[index]
                candidates = np.flatnonzero(scores >= SEMANTIC_SAME_POST_THRESHOLD)
                for position in candidates[np.argsort(scores[candidates])[::-1]]:
                    rep_index = representative_indices[int(position)]
                    rep_row = rows[rep_index]
                    shorter_length = min(
                        len(normalize_text(row["answer"])),
                        len(normalize_text(rep_row["answer"])),
                    )
                    if shorter_length < 12:
                        continue
                    threshold = (
                        SEMANTIC_SAME_POST_THRESHOLD
                        if row["post_url"] == rep_row["post_url"]
                        else SEMANTIC_CROSS_POST_THRESHOLD
                    )
                    score = float(scores[position])
                    if score < threshold:
                        continue
                    if not _semantically_compatible(text, _semantic_text(rep_row)):
                        continue
                    matched_rep_position = int(position)
                    similarity = score
                    reason = "语义相近"
                    break

            if matched_rep_position is None:
                position = representative_count
                representatives[position] = matrix[index]
                representative_indices.append(index)
                representative_count += 1
                exact_representatives.setdefault(row["content_hash"], position)
                assignments[index] = (1, "", 0.0, row["reply_id"], reason)
            else:
                rep_index = representative_indices[matched_rep_position]
                rep_id = rows[rep_index]["reply_id"]
                assignments[index] = (0, rep_id, round(similarity, 6), rep_id, reason)

            if progress and (completed % 256 == 0 or completed == total):
                progress(f"正在合并语义相近回复 {completed}/{total}", completed, total)

        updates = [
            (*assignment, rows[index]["reply_id"])
            for index, assignment in enumerate(assignments)
            if assignment is not None
        ]
        with self.connection:
            self.connection.executemany(
                """
                UPDATE qa_pairs
                SET is_retrievable = ?, duplicate_of = ?, similarity_score = ?,
                    semantic_group_id = ?, dedupe_reason = ?
                WHERE reply_id = ?
                """,
                updates,
            )
            self.connection.execute("DELETE FROM chunks WHERE source_type='qa'")
            representative_rows = self.connection.execute(
                """
                SELECT q.*, p.title AS post_title, p.published_at AS post_published_at
                FROM qa_pairs q JOIN posts p ON p.url = q.post_url
                WHERE q.is_retrievable = 1
                """
            ).fetchall()
            for row in representative_rows:
                question = normalize_text(row["question"])
                answer = normalize_text(row["answer"])
                content = (
                    f"用户问题：{question}\n刺大回复：{answer}"
                    if question else f"刺大评论：{answer}"
                )
                self._insert_chunk(
                    "qa",
                    row["reply_id"],
                    {
                        "url": row["post_url"],
                        "title": row["post_title"],
                        "published_at": row["published_at"] or row["post_published_at"],
                    },
                    content,
                    row["source_url"],
                    0,
                )
            self.connection.execute(
                "INSERT OR REPLACE INTO knowledge_meta(key, value) VALUES('semantic_model', ?)",
                (model_name,),
            )
            self.connection.execute(
                "INSERT OR REPLACE INTO knowledge_meta(key, value) VALUES('semantic_cleaned_at', ?)",
                (cleaned_at,),
            )
            self.connection.execute(
                "DELETE FROM qa_embeddings WHERE reply_id NOT IN (SELECT reply_id FROM qa_pairs)"
            )
        self.rebuild_fts()
        return {
            "original": total,
            "retrievable": representative_count,
            "duplicates": total - representative_count,
            "model": model_name,
            "cleaned_at": cleaned_at,
        }

    def stats(self) -> dict:
        post_count = self.connection.execute("SELECT COUNT(*) FROM posts").fetchone()[0]
        core_count = self.connection.execute(
            "SELECT COUNT(*) FROM posts WHERE scope='top_year'"
        ).fetchone()[0]
        supplemental_count = self.connection.execute(
            "SELECT COUNT(*) FROM posts WHERE scope='recent_qa'"
        ).fetchone()[0]
        archived_count = self.connection.execute(
            "SELECT COUNT(*) FROM posts WHERE scope='recent_archive'"
        ).fetchone()[0]
        qa_count = self.connection.execute("SELECT COUNT(*) FROM qa_pairs").fetchone()[0]
        retrievable_qa_count = self.connection.execute(
            "SELECT COUNT(*) FROM qa_pairs WHERE is_retrievable = 1"
        ).fetchone()[0]
        semantic_duplicate_count = self.connection.execute(
            "SELECT COUNT(*) FROM qa_pairs WHERE is_retrievable = 0"
        ).fetchone()[0]
        community_count = self.connection.execute(
            "SELECT COUNT(*) FROM chunks WHERE source_type='community'"
        ).fetchone()[0]
        manual_source_count = self.connection.execute(
            "SELECT COUNT(*) FROM local_sources"
        ).fetchone()[0]
        manual_chunk_count = self.connection.execute(
            "SELECT COUNT(*) FROM chunks WHERE source_type='manual'"
        ).fetchone()[0]
        chunk_count = self.connection.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
        last_run = self.connection.execute(
            "SELECT finished_at FROM crawl_runs ORDER BY id DESC LIMIT 1"
        ).fetchone()
        semantic_meta = {
            row["key"]: row["value"]
            for row in self.connection.execute(
                "SELECT key, value FROM knowledge_meta WHERE key IN ('semantic_model', 'semantic_cleaned_at')"
            ).fetchall()
        }
        return {
            "posts": post_count,
            "core_posts": core_count,
            "supplemental_posts": supplemental_count,
            "archived_posts": archived_count,
            "qa_pairs": qa_count,
            "retrievable_qa": retrievable_qa_count,
            "semantic_duplicates": semantic_duplicate_count,
            "semantic_model": semantic_meta.get("semantic_model", SEMANTIC_MODEL),
            "semantic_cleaned_at": semantic_meta.get("semantic_cleaned_at", "尚未清洗"),
            "community_comments": community_count,
            "manual_sources": manual_source_count,
            "manual_chunks": manual_chunk_count,
            "chunks": chunk_count,
            "last_sync": last_run["finished_at"] if last_run else "尚未同步",
        }

    def list_posts(self) -> list[dict]:
        return [
            dict(row)
            for row in self.connection.execute(
                """
                SELECT title, published_at, views, reply_count, likes,
                       scope, body_truncated, comments_accessible,
                       useful_comment_count, capture_mode, total_comment_pages,
                       scanned_comment_pages, url
                FROM posts ORDER BY views DESC
                """
            ).fetchall()
        ]

    def record_run(self, result: dict) -> None:
        self.connection.execute(
            """
            INSERT INTO crawl_runs (
                started_at, finished_at, discovered, fetched, reused, failed, error_details
            ) VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                result["started_at"], result["finished_at"], result["discovered"],
                result["fetched"], result["reused"], result["failed"],
                "\n".join(result["errors"]),
            ),
        )
        self.connection.commit()


def sync_knowledge_incremental(
    progress: Callable[[str, int, int], None] | None = None,
) -> dict:
    """首次建立核心 20 篇，之后只增量维护近期 10 篇。"""
    started_at = datetime.now().isoformat(timespec="seconds")
    with TgbCrawler() as crawler, KnowledgeStore() as store:
        core_urls = set(store.scope_urls("top_year"))
        core_posts: list[dict] = []
        needs_core_bootstrap = len(core_urls) < CRAWL_TOP_POST_LIMIT
        if needs_core_bootstrap:
            discovered_core = crawler.discover_top_posts(
                days=365,
                limit=CRAWL_TOP_POST_LIMIT,
                progress=progress,
            )
            core_posts = [
                post for post in discovered_core if post["url"] not in core_urls
            ][: CRAWL_TOP_POST_LIMIT - len(core_urls)]
            core_urls.update(post["url"] for post in core_posts)
        recent_posts = crawler.discover_recent_posts(limit=10)
        supplements = [post for post in recent_posts if post["url"] not in core_urls]
        work_items = [(post, "top_year") for post in core_posts] + [
            (post, "recent_qa") for post in supplements
        ]
        result = {
            "started_at": started_at,
            "finished_at": "",
            "discovered": len(work_items),
            "core_posts": len(core_urls),
            "core_bootstrapped": needs_core_bootstrap,
            "core_frozen": not needs_core_bootstrap,
            "supplemental_posts": len(supplements),
            "archived_this_run": 0,
            "fetched": 0,
            "reused": 0,
            "failed": 0,
            "errors": [],
        }
        for index, (post, scope) in enumerate(work_items, 1):
            if progress:
                label = "高阅读量主帖" if scope == "top_year" else "近期公开问答补充"
                progress(f"正在处理 {index}/{len(work_items)}：{post['title']}（{label}）", index, len(work_items))
            state = store.post_state(post["url"])
            if (
                state
                and state["body_hash"]
                and (
                    state["capture_mode"] == "authenticated_browser"
                    or (
                        state["pipeline_version"] == CRAWL_PIPELINE_VERSION
                        and state["reply_count"] == post["reply_count"]
                    )
                )
            ):
                store.update_post_metrics(post, scope)
                result["reused"] += 1
                continue
            try:
                store.upsert_post(crawler.fetch_post(post), scope=scope)
                result["fetched"] += 1
            except RuntimeError as exc:
                result["failed"] += 1
                result["errors"].append(f"{post['title']}：{exc}")
        if supplements:
            result["archived_this_run"] = store.archive_recent_except(
                [post["url"] for post in supplements]
            )
        result["core_posts"] = len(store.scope_urls("top_year"))
        try:
            result["manual_source"] = store.import_manual_docx()
        except RuntimeError as exc:
            result["manual_source"] = {"imported": False, "error": str(exc)}
        result["semantic_clean"] = store.semantic_clean_qa(progress)
        result["finished_at"] = datetime.now().isoformat(timespec="seconds")
        store.record_run(result)
        return result


def sync_top_year(
    progress: Callable[[str, int, int], None] | None = None,
) -> dict:
    """兼容旧入口；实际执行冻结核心库后的增量更新。"""
    return sync_knowledge_incremental(progress)

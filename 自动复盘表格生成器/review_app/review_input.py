from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path

from docx import Document

from .analysis import workbook_to_text
from .schemas import AnalyzeRequest


MAX_FILE_BYTES = 15 * 1024 * 1024


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ValueError("文本编码无法识别，请保存为 UTF-8 后重试")


def _docx_to_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    lines = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                lines.append(" | ".join(values))
    text = "\n".join(lines).strip()
    if not text:
        raise ValueError("Word 文档中没有可分析的文字")
    return text


def extract_review_text(payload: AnalyzeRequest) -> str:
    """把网页文字或上传文件统一转换为后续清洗需要的纯文本。"""
    if payload.text.strip():
        return payload.text.strip()
    if not payload.content_base64:
        raise ValueError("请先导入复盘文件或自爬取当日复盘")
    try:
        content = base64.b64decode(payload.content_base64, validate=True)
    except ValueError as exc:
        raise ValueError("上传文件内容损坏，请重新选择") from exc
    if len(content) > MAX_FILE_BYTES:
        raise ValueError("复盘文件不能超过 15MB")
    suffix = Path(payload.filename).suffix.lower()
    if suffix in {".txt", ".md"}:
        return _decode_text(content)
    if suffix == ".docx":
        return _docx_to_text(content)
    if suffix == ".xlsx":
        return workbook_to_text(content)
    raise ValueError("仅支持 DOCX、XLSX、TXT 和 Markdown 复盘文件")

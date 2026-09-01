from __future__ import annotations

from array import array
from datetime import date, datetime
from io import BytesIO

import pytest
from docx import Document

from review_app.dragon.analysis import parse_analysis_result
from review_app.dragon.context import build_analysis_context
from review_app.dragon.knowledge import DragonKnowledgeStore, ParsedUnit
from review_app.dragon.market import (
    EastmoneyDragonMarketProvider,
    StaticDragonMarketProvider,
    normalize_market_record,
)
from review_app.dragon.pipeline import DragonPipeline
from review_app.dragon.rules import evaluate_rules
from review_app.dragon.schemas import DragonEvidence, ReviewSnapshot, ReviewSnapshotInput
from review_app.dragon.store import DragonRuntimeStore


TRADE_DATE = date(2026, 8, 28)


def candidate(code: str, name: str, first: str = "09:30", **overrides):
    values = {
        "trade_date": TRADE_DATE,
        "stock_code": code,
        "stock_name": name,
        "first_seal_time": first,
        "last_seal_time": "14:30",
        "break_times": [],
        "board_break_count": 0,
        "final_order_amount": 120_000_000,
        "float_market_cap": 2_000_000_000,
        "pool_board_count": 1,
        "previous_day_limit_up": False,
        "is_confirmed_first_board": True,
        "is_main_board_10cm_ordinary": True,
        "close_limit_up": True,
    }
    values.update(overrides)
    return normalize_market_record(values)


def test_database_paths_use_exact_whitelist(tmp_path):
    DragonRuntimeStore(tmp_path / "dragon_runtime.db")
    with DragonKnowledgeStore(tmp_path / "dragon_knowledge.db"):
        pass
    for name in ("foo.db", "review_jobs.db", "dragon_knowledge.db"):
        with pytest.raises(ValueError):
            DragonRuntimeStore(tmp_path / name)
    for name in ("foo.db", "review_knowledge.db", "dragon_runtime.db"):
        with pytest.raises(ValueError):
            DragonKnowledgeStore(tmp_path / name)


def test_reconfirm_snapshot_updates_json_and_latest_order(tmp_path):
    store = DragonRuntimeStore(tmp_path / "dragon_runtime.db")
    first = store.save_review_snapshot(
        ReviewSnapshotInput(trade_date=TRADE_DATE, market_core="旧结论", confirm_as_layout=True)
    )
    second = store.save_review_snapshot(
        ReviewSnapshotInput(trade_date=TRADE_DATE, market_core="新结论", confirm_as_layout=True)
    )
    snapshots = store.list_review_snapshots(trade_date=TRADE_DATE)
    assert snapshots[0].snapshot_id == second.snapshot_id
    assert snapshots[0].is_confirmed is True
    assert snapshots[1].snapshot_id == first.snapshot_id
    assert snapshots[1].is_confirmed is False
    assert snapshots[1].confirmed_at is None
    assert store.get_review_snapshot(TRADE_DATE).snapshot_id == second.snapshot_id
    assert store.get_review_snapshot(TRADE_DATE, confirmed_only=True).snapshot_id == second.snapshot_id


def test_missing_model_evidence_refs_downgrades_history():
    screening = evaluate_rules(candidate("000001", "甲股"), [])
    snapshot = ReviewSnapshot(
        trade_date=TRADE_DATE,
        market_core="用户确认",
        is_confirmed=True,
        confirm_as_layout=True,
    )
    context = build_analysis_context(
        snapshot=snapshot,
        screening=screening,
        evidence=[DragonEvidence(source_id="doc-1", chunk_id="chunk-1", content="历史案例")],
    )
    result = parse_analysis_result(
        {
            "stock_code": "wrong",
            "stock_name": "wrong",
            "basic_pass": False,
            "conclusion": "观察",
            "historical_models": ["模型A"],
            "historical_recognition": "声称已匹配",
            "evidence_refs": [],
        },
        context,
    )
    assert result.evidence_refs == []
    assert result.historical_models == []
    assert result.historical_recognition == "辨识度不足"


def test_tencent_ohlc_break_timeline_is_stable(monkeypatch):
    provider = EastmoneyDragonMarketProvider()
    payload = {
        "data": {
            "sz000001": {
                "m1": [
                    ["202608281359", "10.00", "10.00", "10.00", "10.00"],
                    ["202608281400", "10.00", "9.80", "10.00", "9.80"],
                    ["202608281401", "9.80", "9.85", "9.90", "9.75"],
                    ["202608281402", "9.85", "10.00", "10.00", "9.80"],
                    ["202608281403", "10.00", "9.90", "10.00", "9.90"],
                ]
            }
        }
    }
    monkeypatch.setattr(provider, "_get_json", lambda *_args, **_kwargs: payload)
    try:
        breaks, source = provider._tencent_break_times(
            "000001", TRADE_DATE, 10.0, "09:30:00"
        )
    finally:
        provider.close()
    assert breaks == ["14:00:00", "14:03:00"]
    assert source == "1m:tencent_ohlc"
    assert provider._clock_from_stamp("2026-08-28 13:59") == "13:59:00"


def test_selected_stock_keeps_full_market_attribute_rank(tmp_path):
    runtime = DragonRuntimeStore(tmp_path / "dragon_runtime.db")
    runtime.ensure_confirmed_default_rules()
    runtime.save_review_snapshot(
        ReviewSnapshotInput(
            trade_date=TRADE_DATE,
            market_core="机器人",
            source_text="机器人板块：甲股、乙股。",
            confirm_as_layout=True,
        )
    )
    provider = StaticDragonMarketProvider(
        [candidate("000001", "甲股", "09:30"), candidate("000002", "乙股", "09:40")]
    )
    with DragonKnowledgeStore(tmp_path / "dragon_knowledge.db") as knowledge:
        output = DragonPipeline(runtime, knowledge, provider).prepare(
            TRADE_DATE, stock_codes=["000002"]
        )
    assert len(output.market_snapshot.candidates) == 2
    assert len(output.screening) == 1
    assert output.screening[0].candidate.same_attribute_orders["机器人"] == 2


def test_empty_review_text_clears_provider_attribute_injection(tmp_path):
    runtime = DragonRuntimeStore(tmp_path / "dragon_runtime.db")
    runtime.ensure_confirmed_default_rules()
    runtime.save_review_snapshot(
        ReviewSnapshotInput(trade_date=TRADE_DATE, market_core="仅结构化结论", confirm_as_layout=True)
    )
    injected = candidate("000001", "甲股").model_copy(
        update={
            "review_attribute_status": "明确匹配",
            "review_attributes": ["行情注入"],
            "same_attribute_orders": {"行情注入": 1},
            "attribute_evidence": [{"source": "provider"}],
        }
    )
    with DragonKnowledgeStore(tmp_path / "dragon_knowledge.db") as knowledge:
        output = DragonPipeline(
            runtime, knowledge, StaticDragonMarketProvider([injected])
        ).prepare(TRADE_DATE)
    normalized = output.screening[0].candidate
    assert normalized.review_attribute_status == "没有提及"
    assert normalized.review_attributes == []
    assert normalized.same_attribute_orders == {}
    assert normalized.attribute_evidence == []


def test_empty_active_rule_version_stops_pipeline(tmp_path):
    runtime = DragonRuntimeStore(tmp_path / "dragon_runtime.db")
    runtime.save_rule_version([], name="空版本", activate=True)
    runtime.save_review_snapshot(
        ReviewSnapshotInput(trade_date=TRADE_DATE, market_core="已确认", confirm_as_layout=True)
    )
    with DragonKnowledgeStore(tmp_path / "dragon_knowledge.db") as knowledge:
        with pytest.raises(ValueError, match="没有启用规则"):
            DragonPipeline(
                runtime,
                knowledge,
                StaticDragonMarketProvider([candidate("000001", "甲股")]),
            ).prepare(TRADE_DATE)


def test_late_break_watch_keeps_other_hard_failures(tmp_path):
    store = DragonRuntimeStore(tmp_path / "dragon_runtime.db")
    version = store.ensure_confirmed_default_rules()
    result = evaluate_rules(
        candidate(
            "000001",
            "甲股",
            break_times=["13:05:00", "13:08:00"],
            board_break_count=2,
            final_order_amount=50_000_000,
        ),
        version.rules,
    )
    assert result.basic_pass is False
    assert result.candidate_bucket == "late_break_watch"
    assert len(result.disqualifying_rule_ids) == 2


def test_first_break_after_1300_drives_late_break_rule_and_under_two_is_relaxed(tmp_path):
    store = DragonRuntimeStore(tmp_path / "dragon_runtime.db")
    version = store.ensure_confirmed_default_rules()

    suspected_only = evaluate_rules(
        candidate(
            "000001",
            "甲股",
            board_break_count=0,
            last_seal_time="14:30",
            break_times=["13:05:00"],
            break_suspected=True,
            break_suspicion_reasons=["公开炸板次数=0，分钟行情识别=1"],
        ),
        version.rules,
    )
    confirmed_late = evaluate_rules(
        candidate(
            "000002",
            "乙股",
            board_break_count=2,
            break_times=["13:05:00", "13:08:00"],
            last_seal_time="13:10",
        ),
        version.rules,
    )
    morning_first = evaluate_rules(
        candidate(
            "000003",
            "丙股",
            board_break_count=2,
            break_times=["10:05:00", "13:08:00"],
            last_seal_time="13:10",
        ),
        version.rules,
    )
    afternoon_first_seal = evaluate_rules(
        candidate(
            "000004",
            "丁股",
            first="13:08:07",
            board_break_count=2,
            break_times=["13:20:00", "14:10:00"],
            last_seal_time="14:25:55",
        ),
        version.rules,
    )

    assert suspected_only.basic_pass is True
    assert suspected_only.candidate.break_suspected is True
    assert suspected_only.candidate.public_late_break is False
    assert confirmed_late.basic_pass is False
    assert confirmed_late.candidate_bucket == "late_break_watch"
    assert morning_first.basic_pass is True
    assert afternoon_first_seal.basic_pass is True
    assert afternoon_first_seal.candidate.public_late_break is False


def test_afternoon_first_seal_does_not_emit_minute_mismatch_suspicion(monkeypatch):
    provider = EastmoneyDragonMarketProvider()
    current_row = {
        "c": "000001", "n": "甲股", "p": 10_000, "lbc": 1,
        "fbt": 130807, "lbt": 142555, "zbc": 2, "fund": 55_545_740,
        "amount": 100_000_000, "hs": 5.0, "ltsz": 9_269_658_706.93,
    }
    monkeypatch.setattr(provider, "_pool", lambda day: [current_row] if day == TRADE_DATE else [])
    monkeypatch.setattr(provider, "_previous_trade_date", lambda _day: TRADE_DATE.replace(day=27))
    monkeypatch.setattr(provider, "_trend_break_times", lambda *_args, **_kwargs: (["13:20:00"], "1m:fixture"))
    try:
        found = provider.fetch_first_board_candidates(TRADE_DATE)
    finally:
        provider.close()

    assert len(found) == 1
    assert found[0].public_late_break is False
    assert found[0].break_suspected is False
    assert found[0].break_suspicion_reasons == []


def test_docx_case_headings_build_exact_name_route(tmp_path):
    document = Document()
    document.add_paragraph("短线记录 654321")
    document.add_paragraph("2026年1月")
    document.add_paragraph("创新医疗 1/05-1/09 5板")
    document.add_paragraph("1/05 首板原因：主动走强。")
    document.add_paragraph("银河电子 1/05-1/12 6板")
    document.add_paragraph("1/12 最高板复盘：完成穿越。")
    stream = BytesIO()
    document.save(stream)
    with DragonKnowledgeStore(tmp_path / "dragon_knowledge.db") as store:
        store.import_bytes(
            "历史高标.docx",
            stream.getvalue(),
            tags=["历史高标"],
            build_semantic=False,
        )
        innovation = store.search("创新医疗", stock_name="创新医疗", semantic=False)
        fake = store.search("654321", stock_code="654321", semantic=False)
    assert innovation and innovation[0]["exact_score"] > 0
    assert innovation[0]["stock_name"] == "创新医疗"
    assert all(item["exact_score"] == 0 for item in fake)


def test_semantic_status_reports_partial_index(tmp_path):
    with DragonKnowledgeStore(tmp_path / "dragon_knowledge.db") as store:
        store.import_text(
            ("第一段历史案例。" * 220) + "\n\n" + ("第二段历史案例。" * 220),
            source_path="inline://partial",
            build_semantic=False,
        )
        row = store.connection.execute(
            "SELECT id, content_hash FROM chunks ORDER BY id LIMIT 1"
        ).fetchone()
        with store.connection:
            store.connection.execute(
                """
                INSERT INTO embeddings(chunk_id, model, content_hash, dimensions, vector, updated_at)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (
                    row["id"],
                    store.semantic_model,
                    row["content_hash"],
                    1,
                    array("f", [1.0]).tobytes(),
                    datetime.now().isoformat(),
                ),
            )
        status = store.semantic_status()
    assert status["status"] == "partial"
    assert status["current_embeddings"] == 1
    assert status["total_chunks"] > 1


def test_reimport_reuses_unchanged_chunk_embeddings(tmp_path):
    embedded_batches: list[list[str]] = []

    with DragonKnowledgeStore(tmp_path / "dragon_knowledge.db") as store:
        store._embed = lambda texts: (  # type: ignore[method-assign]
            embedded_batches.append(list(texts))
            or [array("f", [float(index + 1)]) for index, _ in enumerate(texts)]
        )

        def import_units(contents: list[str]):
            return store.import_text(
                "\n\n".join(contents),
                source_path="inline://cache-regression",
                units=[ParsedUnit(content) for content in contents],
            )

        first = import_units(["第一段历史案例", "第二段历史案例", "第三段历史案例"])
        first_rows = store.connection.execute(
            "SELECT id, chunk_index, content_hash FROM chunks ORDER BY chunk_index"
        ).fetchall()
        assert first["semantic"]["status"] == "ready"
        assert first["semantic"]["indexed"] == 3
        assert first["semantic"]["cached"] == 0
        assert first["semantic"]["total"] == 3
        assert first["semantic"]["cache_hit_rate"] == 0.0
        assert [len(batch) for batch in embedded_batches] == [3]

        repeated = import_units(["第一段历史案例", "第二段历史案例", "第三段历史案例"])
        repeated_rows = store.connection.execute(
            "SELECT id, chunk_index, content_hash FROM chunks ORDER BY chunk_index"
        ).fetchall()
        assert repeated["semantic"]["indexed"] == 0
        assert repeated["semantic"]["cached"] == 3
        assert repeated["semantic"]["total"] == 3
        assert repeated["semantic"]["cache_hit_rate"] == 1.0
        assert [row["id"] for row in repeated_rows] == [row["id"] for row in first_rows]
        assert [len(batch) for batch in embedded_batches] == [3]

        changed = import_units(["第一段历史案例", "第二段历史案例（修订）", "第三段历史案例"])
        changed_rows = store.connection.execute(
            "SELECT id, chunk_index, content_hash FROM chunks ORDER BY chunk_index"
        ).fetchall()
        assert changed["semantic"]["indexed"] == 1
        assert changed["semantic"]["cached"] == 2
        assert changed["semantic"]["total"] == 3
        assert [row["id"] for row in changed_rows] == [row["id"] for row in first_rows]
        assert [len(batch) for batch in embedded_batches] == [3, 1]

        shortened = import_units(["第一段历史案例", "第二段历史案例（修订）"])
        remaining_rows = store.connection.execute(
            "SELECT id FROM chunks ORDER BY chunk_index"
        ).fetchall()
        embedding_count = store.connection.execute(
            "SELECT COUNT(*) FROM embeddings"
        ).fetchone()[0]
        assert shortened["semantic"]["indexed"] == 0
        assert shortened["semantic"]["cached"] == 2
        assert shortened["semantic"]["total"] == 2
        assert [row["id"] for row in remaining_rows] == [row["id"] for row in first_rows[:2]]
        assert embedding_count == 2

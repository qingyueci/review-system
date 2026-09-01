from io import BytesIO

from docx import Document
import numpy as np

import review_app.knowledge as knowledge_module
from review_app.analysis import ANALYSIS_SYSTEM_PROMPT
from review_app.cleaning import is_greeting_or_noise, normalize_text
from review_app.crawler import TgbCrawler, _sample_pages
from review_app.docx_export import generate_analysis_docx
from review_app.knowledge import KnowledgeStore


def _knowledge_post(url: str, title: str, body: str, reply_count: int = 1) -> dict:
    return {
        "url": url,
        "topic_id": url.rsplit("/", 1)[-1],
        "title": title,
        "published_at": "2026-07-18T16:20",
        "views": 100,
        "reply_count": reply_count,
        "likes": 5,
        "summary": title,
        "body": body,
        "body_hash": f"hash-{title}",
        "total_comment_pages": 1,
        "scanned_comment_pages": [1],
        "author_replies": [],
        "community_comments": [],
    }


def test_cleaning_filters_greetings_but_keeps_layout_questions():
    assert is_greeting_or_noise("先赞后看，刺大发财！")
    assert not is_greeting_or_noise("恒尚首板出身决定了它今天应该完成什么任务？")
    assert normalize_text("正文\u200b\n\n\n下载淘股吧APP") == "正文"


def test_sample_comment_pages_covers_beginning_middle_and_end():
    pages = _sample_pages(39, 12)
    assert len(pages) == 12
    assert pages[:3] == [1, 2, 3]
    assert pages[-3:] == [37, 38, 39]
    assert any(10 < page < 30 for page in pages)


def test_mobile_listing_parser_reads_metrics():
    html = """
    <div class="indexContentItem">
      <span contentid="123" subject="摘要"></span>
      <span class="content_time"><span>2026-07-18 16:20</span></span>
      <a class="contentTitle" href="/a/demo">718</a>
      <a class="content_text">[摘要] 今日复盘</a>
      <div class="viewBtn"><span>浏览(123,456)</span></div>
      <div class="plBtn"><span>评论(789)</span></div>
      <div class="zanBtn"><span>赞(66)</span></div>
    </div>
    """
    posts = TgbCrawler._parse_listing(html)
    assert posts[0]["topic_id"] == "123"
    assert posts[0]["views"] == 123456
    assert posts[0]["reply_count"] == 789
    assert posts[0]["url"] == "https://www.tgb.cn/a/demo"


def test_community_comments_use_likes_and_layout_relevance():
    comments = [
        {"reply_id": "1", "answer": "这只票首板出身决定任务，今天主动性也完成了板块反推。", "likes": 2},
        {"reply_id": "2", "answer": "写得不错，感谢分享，继续学习，祝老师天天发财。", "likes": 12},
        {"reply_id": "3", "answer": "首板出身、任务、地位、主动性和板块协同都应该放在一起理解。", "likes": 1},
        {"reply_id": "4", "answer": "普通长评论但没有足够点赞，也没有布局相关信息。", "likes": 0},
    ]
    selected = TgbCrawler._select_community_comments(comments)
    assert [item["reply_id"] for item in selected] == ["1", "3"]


def test_useful_comment_parser_keeps_author_and_quote_context():
    post = {
        "url": "https://www.tgb.cn/a/test",
        "title": "测试",
        "published_at": "2026-07-18T16:00",
    }
    item = TgbCrawler._parse_useful_comment({
        "replyID": 123,
        "userID": 5894557,
        "userName": "延边刺客",
        "body": "首板出身决定了它的任务。<br/>不是只看报价。",
        "quoteContent": "它今天承担什么任务？",
        "quoteUserID": 8,
        "quoteUserName": "提问者",
        "usefulNum": 66,
        "pageNo": 5,
    }, post)
    assert item["reply_id"] == "123"
    assert item["author_id"] == "5894557"
    assert "首板出身" in item["answer"]
    assert item["question_author"] == "提问者"
    assert item["likes"] == 66
    assert item["source_url"].endswith("-5#reply123")


def test_knowledge_store_keeps_question_and_answer_together(tmp_path):
    post = {
        "url": "https://www.tgb.cn/a/test",
        "topic_id": "1",
        "title": "测试复盘",
        "published_at": "2026-07-18T16:20",
        "views": 100,
        "reply_count": 10,
        "likes": 5,
        "summary": "摘要",
        "body": "机器人首板出身，任务是为科技方向做发酵确认。",
        "body_hash": "hash",
        "total_comment_pages": 1,
        "scanned_comment_pages": [1],
        "author_replies": [{
            "reply_id": "99",
            "question": "这只股票今天在布局里承担什么任务？",
            "answer": "先看首板出身，它的任务是反推机器人方向，而不是只看报价。",
            "question_author": "测试用户",
            "published_at": "2026-07-18 18:00",
            "floor": 20,
            "likes": 8,
            "source_url": "https://www.tgb.cn/a/test-1#reply99",
        }],
        "community_comments": [{
            "reply_id": "100",
            "author_name": "社区用户",
            "question": "",
            "answer": "这个首板出身对应的是板块发酵任务，需要观察主动性。",
            "published_at": "2026-07-18 18:10",
            "floor": 21,
            "likes": 10,
            "source_url": "https://www.tgb.cn/a/test-1#reply100",
        }],
    }
    with KnowledgeStore(tmp_path / "knowledge.db") as store:
        store.upsert_post(post)
        results = store.search("首板出身 个股任务 机器人", limit=5)
        assert results
        assert results[0]["retrieval_score"] > 0
        assert "本地向量" in results[0]["retrieval_mode"]
        combined = "\n".join(item["content"] for item in results)
        assert "用户问题" in combined
        assert "刺大回复" in combined
        assert store.stats()["community_comments"] == 1


def test_hybrid_search_uses_vector_concepts_and_source_weight(tmp_path):
    with KnowledgeStore(tmp_path / "knowledge.db") as store:
        post = {
            "url": "https://example.com/source",
            "title": "布局讨论",
            "published_at": "2026-07-18T10:00",
        }
        content = "该股承担板块角色，需要观察协同和带动。"
        store._insert_chunk(
            "qa", "qa-1", post, content, post["url"], 0,
        )
        store._insert_chunk(
            "community", "community-1", post, content, post["url"] + "#2", 0,
        )
        store.rebuild_fts()

        results = store.search("这只股票的使命是什么", limit=5)

        assert results
        assert results[0]["source_type"] == "qa"
        assert results[0]["vector_score"] > 0
        assert results[0]["source_weight"] == 1.0
        community = next(
            item for item in results if item["source_type"] == "community"
        )
        assert results[0]["retrieval_score"] > community["retrieval_score"]


def test_semantic_clean_keeps_raw_replies_and_indexes_representatives(
    tmp_path, monkeypatch
):
    post = _knowledge_post(
        "https://www.tgb.cn/a/semantic", "语义清洗", "首板出身决定个股任务。", 3
    )
    post["author_replies"] = [
        {
            "reply_id": "qa-1", "question": "怎么看首板",
            "answer": "首板的发酵来源决定了个股后续承担的任务，需要结合板块主动性判断。",
            "published_at": "2026-07-18 18:00", "floor": 1, "likes": 8,
            "source_url": "https://www.tgb.cn/a/semantic#1",
        },
        {
            "reply_id": "qa-2", "question": "怎么看个股任务",
            "answer": "个股后续承担什么任务，首先要回到首板的发酵来源和板块主动性。",
            "published_at": "2026-07-18 18:01", "floor": 2, "likes": 2,
            "source_url": "https://www.tgb.cn/a/semantic#2",
        },
        {
            "reply_id": "qa-3", "question": "市场走弱怎么办",
            "answer": "市场整体走弱时先降低仓位，等待新的确认信号，不要急着扩大交易。",
            "published_at": "2026-07-18 18:02", "floor": 3, "likes": 1,
            "source_url": "https://www.tgb.cn/a/semantic#3",
        },
    ]

    class FakeEmbedding:
        def embed(self, texts, batch_size=64):
            assert batch_size == 64
            vectors = {
                post["author_replies"][0]["answer"]: np.array([1.0, 0.0, 0.0]),
                post["author_replies"][1]["answer"]: np.array([0.999, 0.01, 0.0]),
                post["author_replies"][2]["answer"]: np.array([0.0, 1.0, 0.0]),
            }
            return (vectors[text] for text in texts)

    monkeypatch.setattr(knowledge_module, "_embedding_model", lambda *_args: FakeEmbedding())
    with KnowledgeStore(tmp_path / "knowledge.db") as store:
        store.upsert_post(post)
        cleaned = store.semantic_clean_qa()
        stats = store.stats()
        raw_count = store.connection.execute("SELECT COUNT(*) FROM qa_pairs").fetchone()[0]
        indexed_count = store.connection.execute(
            "SELECT COUNT(*) FROM chunks WHERE source_type='qa'"
        ).fetchone()[0]

    assert cleaned["original"] == raw_count == 3
    assert cleaned["retrievable"] == indexed_count == 2
    assert cleaned["duplicates"] == stats["semantic_duplicates"] == 1


def test_incremental_sync_freezes_core_and_archives_old_recent(tmp_path, monkeypatch):
    database_path = tmp_path / "knowledge.db"
    core = _knowledge_post(
        "https://www.tgb.cn/a/core", "固定核心", "首板出身决定核心任务。"
    )
    old_recent = _knowledge_post(
        "https://www.tgb.cn/a/old", "旧近期帖", "历史归档仍有布局参考价值。"
    )
    with KnowledgeStore(database_path) as store:
        store.upsert_post(core, scope="top_year")
        store.upsert_post(old_recent, scope="recent_qa")

    new_listing = {
        "url": "https://www.tgb.cn/a/new",
        "topic_id": "new",
        "title": "最新复盘",
        "published_at": "2026-07-22T16:20",
        "views": 200,
        "reply_count": 2,
        "likes": 8,
        "summary": "最新复盘",
    }

    class TestStore(KnowledgeStore):
        def __init__(self):
            super().__init__(database_path)

    class FakeCrawler:
        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return None

        def discover_top_posts(self, **_kwargs):
            raise AssertionError("已有核心库时不应重新扫描前 20 篇")

        def discover_recent_posts(self, *, limit):
            assert limit == 10
            return [new_listing]

        def fetch_post(self, post):
            return _knowledge_post(
                post["url"], post["title"], "最新个股承担板块发酵任务。", post["reply_count"]
            )

    monkeypatch.setattr(knowledge_module, "CRAWL_TOP_POST_LIMIT", 1)
    monkeypatch.setattr(knowledge_module, "KnowledgeStore", TestStore)
    monkeypatch.setattr(knowledge_module, "TgbCrawler", FakeCrawler)

    result = knowledge_module.sync_knowledge_incremental()

    assert result["core_frozen"] is True
    assert result["fetched"] == 1
    assert result["archived_this_run"] == 1
    with KnowledgeStore(database_path) as store:
        assert store.scope_urls("top_year") == [core["url"]]
        assert store.scope_urls("recent_qa") == [new_listing["url"]]
        assert store.scope_urls("recent_archive") == [old_recent["url"]]
        archived = store.search("历史归档 布局参考价值", limit=5)
        assert any(item["post_url"] == old_recent["url"] for item in archived)
        archived_source = next(
            item for item in archived if item["post_url"] == old_recent["url"]
        )
        assert archived_source["source_weight"] == 0.648


def test_manual_system_docx_is_imported_as_rag_source(tmp_path):
    source_path = tmp_path / "延边刺客短线打板体系.docx"
    document = Document()
    document.add_heading("首板出身与任务", level=1)
    document.add_paragraph("首板出身决定初始地位，个股需要完成板块发酵和反推任务。")
    document.save(source_path)
    with KnowledgeStore(tmp_path / "knowledge.db") as store:
        result = store.import_manual_docx(source_path)
        sources = store.search("首板出身 初始地位 反推任务", limit=5)
        assert result["imported"]
        assert result["chunks"] == 1
        assert sources[0]["source_type"] == "manual"
        assert store.stats()["manual_sources"] == 1


def test_analysis_prompt_prioritizes_tasks_over_technical_indicators():
    assert "首板出身 -> 原始任务" in ANALYSIS_SYSTEM_PROMPT
    assert "禁止把普通技术分析写成主线" in ANALYSIS_SYSTEM_PROMPT
    assert "个股不是孤立报价" in ANALYSIS_SYSTEM_PROMPT
    assert "社区精选评论" in ANALYSIS_SYSTEM_PROMPT
    assert "人工整理体系" in ANALYSIS_SYSTEM_PROMPT


def test_analysis_docx_contains_sources_and_disclaimer():
    content, filename = generate_analysis_docx(
        "# 今日核心判断\n个股任务优先。\n## 个股任务表\n- 测试股：完成发酵任务",
        [{
            "title": "历史复盘",
            "published_at": "2026-07-01T16:20",
            "source_type": "qa",
            "source_url": "https://www.tgb.cn/a/test",
        }],
        review_date="2026-07-18",
    )
    document = Document(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "不代表原作者本人观点" in text
    assert "个股任务优先" in text
    assert filename == "刺大框架复盘分析_2026-07-18.docx"

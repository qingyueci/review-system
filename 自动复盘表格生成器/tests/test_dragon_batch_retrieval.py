from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document
import pytest

from review_app.dragon.analysis import (
    BATCH_ANALYSIS_SYSTEM_PROMPT,
    DragonAnalysisService,
    DragonBatchValidationError,
    parse_batch_analysis_result,
)
from review_app.dragon.context import build_analysis_context
from review_app.dragon.knowledge import DragonKnowledgeStore
from review_app.dragon.market import normalize_market_record
from review_app.dragon.retrieval import DragonRetriever
from review_app.dragon.rules import evaluate_rules
from review_app.dragon.schemas import DragonSelectionPolicy, ReviewSnapshot


TRADE_DATE = date(2026, 8, 28)


def _candidate(code: str, name: str):
    return normalize_market_record(
        {
            "trade_date": TRADE_DATE,
            "stock_code": code,
            "stock_name": name,
            "first_seal_time": "09:30",
            "is_confirmed_first_board": True,
        }
    )


def _snapshot():
    return ReviewSnapshot(
        trade_date=TRADE_DATE,
        market_core="消费",
        source_text="消费\n乙股\n板块一点响应也没有",
        confirm_as_layout=True,
        is_confirmed=True,
    )


def test_case_layer_keeps_two_stocks_in_one_doc_and_expands_parent(tmp_path):
    document = Document()
    document.add_paragraph("2026年1月")
    document.add_paragraph("甲股 1/05-1/09 5板")
    document.add_paragraph("1/06 甲股主动走强，正文提及乙股但不改变案例归属。")
    document.add_paragraph("乙股 1/10-1/16 7板")
    document.add_paragraph("1/16 乙股完成穿越。")
    stream = BytesIO()
    document.save(stream)

    with DragonKnowledgeStore(tmp_path / "dragon_knowledge.db") as store:
        store.import_bytes("两股历史高标.docx", stream.getvalue(), build_semantic=False)
        assert store.stats()["cases"] == 2
        cases = store.connection.execute(
            "SELECT stock_name, start_date, end_date FROM cases ORDER BY case_index"
        ).fetchall()
        assert [(row["stock_name"], row["start_date"], row["end_date"]) for row in cases] == [
            ("甲股", "2026-01-05", "2026-01-09"),
            ("乙股", "2026-01-10", "2026-01-16"),
        ]
        result = DragonRetriever(store).retrieve(_candidate("000002", "乙股"), _snapshot())

    assert result.evidence_card["history"]["highest_board"] == 7
    assert result.evidence_card["history"]["auto_old_dragon"] is False
    assert result.evidence_card["history"]["dates"] == ["2026-01-10—2026-01-16"]
    assert any("完成穿越" in item.content for item in result.evidence)


def test_batch_decision_calls_model_once_and_keeps_free_analysis():
    snapshot = _snapshot()
    contexts = [
        build_analysis_context(snapshot=snapshot, screening=evaluate_rules(candidate, []))
        for candidate in (_candidate("000001", "甲股"), _candidate("000002", "乙股"))
    ]
    calls = []

    def completion(request):
        calls.append(request.user_prompt)
        return {
            "batch_summary": {"结论": "乙股相对更优"},
            "results": [
                {
                    "stock_code": "000001",
                    "conclusion": "观察",
                    "review_conflict": False,
                    "history_dates": [],
                    "analysis": {"原因": "等待方向响应"},
                },
                {
                    "stock_code": "000002",
                    "conclusion": "重点",
                    "review_conflict": False,
                    "history_dates": [],
                    "analysis": {"模型自定结构": {"任务": "次日验证", "条件": ["主动走强"]}},
                },
            ],
        }

    outcome = DragonAnalysisService(completion).analyze_batch(
        contexts, policy=DragonSelectionPolicy(), model="fixture"
    )
    assert len(calls) == 1
    assert [item.conclusion for item in outcome.results] == ["观察", "重点"]
    assert outcome.results[1].analysis["模型自定结构"]["任务"] == "次日验证"


def test_invalid_zero_focus_is_corrected_once():
    snapshot = _snapshot()
    contexts = [
        build_analysis_context(snapshot=snapshot, screening=evaluate_rules(_candidate("000001", "甲股"), []))
    ]
    calls = []

    def completion(request):
        calls.append(request.user_prompt)
        if len(calls) == 1:
            return {
                "zero_focus_reason": "没有完美候选",
                "results": [{
                    "stock_code": "000001", "conclusion": "观察",
                    "review_conflict": False, "history_dates": [], "analysis": {},
                }],
            }
        return {
            "results": [{
                "stock_code": "000001", "conclusion": "重点",
                "review_conflict": False, "history_dates": [], "analysis": {"修正": "相对第一"},
            }],
        }

    outcome = DragonAnalysisService(completion).analyze_batch(
        contexts, policy=DragonSelectionPolicy(), model="fixture"
    )
    assert len(calls) == 2
    assert outcome.audit["retry_count"] == 1
    assert outcome.results[0].conclusion == "重点"


@pytest.mark.parametrize(
    "reason",
    [
        "与用户确认复盘方向冲突",
        "与用户确认复盘方向冲突：具体说明",
        "与用户确认复盘方向冲突（具体说明）",
        "与用户确认复盘方向冲突(具体说明)",
    ],
)
def test_qualified_exclusion_accepts_fixed_reason_with_explanation(reason):
    context = build_analysis_context(
        snapshot=_snapshot(),
        screening=evaluate_rules(_candidate("000001", "甲股"), []),
    )
    results, _, _ = parse_batch_analysis_result(
        {
            "zero_focus_reason": "与用户确认复盘方向冲突",
            "results": [{
                "stock_code": "000001",
                "conclusion": "排除",
                "review_conflict": True,
                "exclusion_reason": reason,
                "history_dates": [],
                "analysis": {},
            }]
        },
        [context],
        policy=DragonSelectionPolicy(),
    )
    assert results[0].conclusion == "排除"


def test_qualified_exclusion_still_rejects_other_reason():
    context = build_analysis_context(
        snapshot=_snapshot(),
        screening=evaluate_rules(_candidate("000001", "甲股"), []),
    )
    with pytest.raises(DragonBatchValidationError, match="基础合格候选排除理由不被允许"):
        parse_batch_analysis_result(
            {
                "results": [{
                    "stock_code": "000001",
                    "conclusion": "排除",
                    "review_conflict": True,
                    "exclusion_reason": "板块没有响应",
                    "history_dates": [],
                    "analysis": {},
                }]
            },
            [context],
            policy=DragonSelectionPolicy(),
        )


@pytest.mark.parametrize(
    "reason",
    [
        "与用户确认复盘方向冲突：传媒应用让位算力",
        "与用户确认复盘方向冲突：该方向不是主线",
        "与用户确认复盘方向冲突：仅为主线狗腿子",
        "与用户确认复盘方向冲突：板块无响应、无发酵",
    ],
)
def test_relative_priority_language_cannot_exclude_qualified_candidate(reason):
    context = build_analysis_context(
        snapshot=_snapshot(),
        screening=evaluate_rules(_candidate("000001", "甲股"), []),
    )
    with pytest.raises(DragonBatchValidationError, match="基础合格候选排除理由不被允许"):
        parse_batch_analysis_result(
            {
                "results": [{
                    "stock_code": "000001",
                    "conclusion": "排除",
                    "review_conflict": True,
                    "exclusion_reason": reason,
                    "history_dates": [],
                    "analysis": {},
                }]
            },
            [context],
            policy=DragonSelectionPolicy(),
        )


def test_batch_prompt_distinguishes_board_order_from_board_height():
    assert "same_attribute_orders 仅表示同属性首板的日内上板先后" in BATCH_ANALYSIS_SYSTEM_PROMPT
    assert "不得把顺序靠后写成“身位靠后”" in BATCH_ANALYSIS_SYSTEM_PROMPT
